#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""改进版论文生成器 - Part1: 格式辅助函数 + 封面/摘要/目录
格式依据：附件3：江西农业大学本科毕业论文模板.doc
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = Path(__file__).resolve().parent
FIGURES = '/home/pengyiju/code/lunwen/lunwen/6020222035-彭益举/figures/'
OUT = '/home/pengyiju/code/lunwen/lunwen/6020222035-彭益举/2425_41_10475_080902_6020222035_LW.docx'
FRONT_MATTER_TEMPLATE = next(BASE_DIR.rglob('2425_41_10475_080902_602022203.docx'))

doc = Document(str(FRONT_MATTER_TEMPLATE))
doc.core_properties.author = '彭益举'
doc.core_properties.last_modified_by = '彭益举'
doc.core_properties.title = '基于微服务架构的高校智能选课系统设计与实现'
doc.core_properties.subject = '江西农业大学本科毕业论文'

BODY_HEADER = '江西农业大学本科毕业论文（设计、创作）'

COVER_REPLACEMENTS = {
    '学  院：          计算机与信息工程学院': '学  院：          软件学院',
    '班  级：         软件工程2022-1班': '班  级：         2206',
}

# ── 页面设置：A4，与范本一致 ────────────────────────────────────
def apply_section_layout(section):
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)


def clear_story(story):
    p = story.paragraphs[0] if story.paragraphs else story.add_paragraph()
    p.clear()
    pPr = p._p.get_or_add_pPr()
    existing_pbdr = pPr.find(qn('w:pBdr'))
    if existing_pbdr is not None:
        pPr.remove(existing_pbdr)
    for old_p in story.paragraphs[1:]:
        old_p._element.getparent().remove(old_p._element)
    return p


def set_page_numbering(section, fmt=None, start=None):
    sectPr = section._sectPr
    existing = sectPr.find(qn('w:pgNumType'))
    if existing is not None:
        sectPr.remove(existing)
    if fmt is None and start is None:
        return

    pg_num = OxmlElement('w:pgNumType')
    if fmt is not None:
        pg_num.set(qn('w:fmt'), fmt)
    if start is not None:
        pg_num.set(qn('w:start'), str(start))

    insert_at = len(sectPr)
    for idx, child in enumerate(sectPr):
        if child.tag in {qn('w:cols'), qn('w:docGrid')}:
            insert_at = idx
            break
    sectPr.insert(insert_at, pg_num)


def set_header_text(paragraph, text):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(0)

    pPr = paragraph._p.get_or_add_pPr()
    existing_pbdr = pPr.find(qn('w:pBdr'))
    if existing_pbdr is not None:
        pPr.remove(existing_pbdr)
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pbdr.append(bottom)
    pPr.append(pbdr)

    run = paragraph.add_run(text)
    ef(run, east='宋体', west='Times New Roman', size=10.5, bold=False)


def set_footer_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(0)

    run = paragraph.add_run()
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10.5)
    set_east_asia_font(run, '宋体')
    for tag, val in [('w:fldChar', 'begin'), ('w:instrText', 'PAGE'), ('w:fldChar', 'end')]:
        el = OxmlElement(tag)
        if tag == 'w:instrText':
            el.text = val
        else:
            el.set(qn('w:fldCharType'), val)
        run._r.append(el)


def configure_section(section, header_text=None, show_page_number=False,
                      page_number_format=None, page_number_start=None):
    apply_section_layout(section)
    section.different_first_page_header_footer = False
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    header_p = clear_story(section.header)
    footer_p = clear_story(section.footer)
    if header_text:
        set_header_text(header_p, header_text)
    if show_page_number:
        set_footer_page_number(footer_p)
    set_page_numbering(section, page_number_format, page_number_start)


def start_section(header_text=None, show_page_number=False,
                  page_number_format=None, page_number_start=None):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(
        section,
        header_text=header_text,
        show_page_number=show_page_number,
        page_number_format=page_number_format,
        page_number_start=page_number_start,
    )
    return section


def trim_document_to_prefix(keep_body_elements):
    """保留旧文档前 keep_body_elements 个 body 元素，删除其后的旧正文。"""
    body = doc._body._element
    children = list(body)
    for child in children[keep_body_elements:]:
        if child.tag == qn('w:sectPr'):
            continue
        body.remove(child)


def ensure_page_break_before_text(target_text):
    """确保指定段落从新页开始。"""
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == target_text:
            prev_para = paragraph._p.getprevious()
            if prev_para is not None and ''.join(prev_para.itertext()).strip() == '':
                has_page_break = False
                for br in prev_para.iter(qn('w:br')):
                    if br.get(qn('w:type')) == 'page':
                        has_page_break = True
                        break
                if not has_page_break:
                    prev_para.getparent().remove(prev_para)
            prev = paragraph._p.getprevious()
            if prev is not None:
                for br in prev.iter(qn('w:br')):
                    if br.get(qn('w:type')) == 'page':
                        return
            breaker = paragraph.insert_paragraph_before('')
            breaker.add_run().add_break(WD_BREAK.PAGE)
            return


trim_document_to_prefix(26)
ensure_page_break_before_text('江西农业大学')
configure_section(doc.sections[0], header_text=None, show_page_number=False)


def apply_cover_replacements():
    for paragraph in doc.paragraphs[:20]:
        full_text = ''.join(run.text for run in paragraph.runs)
        replacement = COVER_REPLACEMENTS.get(full_text)
        if replacement is None:
            continue
        if paragraph.runs:
            target_index = 0
            for idx, run in enumerate(paragraph.runs):
                if run.text:
                    target_index = idx
                    break
            paragraph.runs[target_index].text = replacement
            for idx, run in enumerate(paragraph.runs):
                # Keep non-text runs intact because the cover underline is
                # stored as a drawing object in the first run.
                if idx != target_index and run.text:
                    run.text = ''
        else:
            paragraph.add_run(replacement)

    # Preserve the original cover layout for class line by keeping the value
    # in the middle run instead of collapsing the whole line into one run.
    class_para = doc.paragraphs[14]
    if len(class_para.runs) >= 4:
        class_para.runs[1].text = '班  级：         '
        class_para.runs[2].text = '2206'
        class_para.runs[3].text = ''


apply_cover_replacements()

# ══════════════════════════════════════════════════════════════
# 辅助函数（严格按模板格式）
# ══════════════════════════════════════════════════════════════

def set_east_asia_font(run, name):
    """同时设置中西文字体"""
    rpr = run._r.get_or_add_rPr()
    rFonts = rpr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), name)


def set_west_font(run, name):
    rpr = run._r.get_or_add_rPr()
    rFonts = rpr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)


def set_bool_prop(run, tag, enabled):
    rpr = run._r.get_or_add_rPr()
    el = rpr.find(qn(tag))
    if el is None:
        el = OxmlElement(tag)
        rpr.append(el)
    el.set(qn('w:val'), '1' if enabled else '0')


def ef(run, east='宋体', west=None, size=12, bold=False):
    """设置run的字体：中文字体(east)、西文字体(west)、字号、粗体"""
    if east:
        set_east_asia_font(run, east)
    if west:
        run.font.name = west
        set_west_font(run, west)
    elif east:
        run.font.name = east
    run.font.size = Pt(size)
    run.font.bold = bold
    set_bool_prop(run, 'w:b', bold)
    set_bool_prop(run, 'w:bCs', bold)


def get_or_add_style(name, base='Normal'):
    styles = doc.styles
    try:
        style = styles[name]
    except KeyError:
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    if base and style.base_style != styles[base]:
        style.base_style = styles[base]
    return style


def set_style_font(style, east='宋体', west='Times New Roman', size=12, bold=False):
    style.font.name = west or east
    style.font.size = Pt(size)
    style.font.bold = bold
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), east)
    if west:
        rFonts.set(qn('w:ascii'), west)
        rFonts.set(qn('w:hAnsi'), west)
        rFonts.set(qn('w:cs'), west)


def configure_styles():
    normal = doc.styles['Normal']
    set_style_font(normal, east='宋体', west='Times New Roman', size=12, bold=False)

    heading1 = get_or_add_style('Thesis Heading 1')
    set_style_font(heading1, east='黑体', west='Times New Roman', size=14, bold=False)
    heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading1.paragraph_format.space_before = Pt(12)
    heading1.paragraph_format.space_after = Pt(12)
    heading1.paragraph_format.line_spacing = Pt(20)
    heading1.paragraph_format.first_line_indent = Pt(0)
    heading1.paragraph_format.left_indent = Pt(0)

    heading2 = get_or_add_style('Thesis Heading 2')
    set_style_font(heading2, east='黑体', west='Times New Roman', size=12, bold=False)
    heading2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading2.paragraph_format.space_before = Pt(6)
    heading2.paragraph_format.space_after = Pt(3)
    heading2.paragraph_format.line_spacing = Pt(20)
    heading2.paragraph_format.first_line_indent = Pt(0)
    heading2.paragraph_format.left_indent = Pt(0)

    heading3 = get_or_add_style('Thesis Heading 3')
    set_style_font(heading3, east='黑体', west='Times New Roman', size=12, bold=False)
    heading3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading3.paragraph_format.space_before = Pt(3)
    heading3.paragraph_format.space_after = Pt(0)
    heading3.paragraph_format.line_spacing = Pt(20)
    heading3.paragraph_format.first_line_indent = Pt(0)
    heading3.paragraph_format.left_indent = Pt(0)

    body_style = get_or_add_style('Thesis Body')
    set_style_font(body_style, east='宋体', west='Times New Roman', size=12, bold=False)
    body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body_style.paragraph_format.space_before = Pt(0)
    body_style.paragraph_format.space_after = Pt(0)
    body_style.paragraph_format.line_spacing = Pt(20)
    body_style.paragraph_format.first_line_indent = Pt(24)

    en_body_style = get_or_add_style('Thesis English Body')
    set_style_font(en_body_style, east='Times New Roman', west='Times New Roman', size=12, bold=False)
    en_body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    en_body_style.paragraph_format.space_before = Pt(0)
    en_body_style.paragraph_format.space_after = Pt(0)
    en_body_style.paragraph_format.line_spacing = Pt(20)
    en_body_style.paragraph_format.first_line_indent = Pt(24)

    center_title = get_or_add_style('Thesis Center Title')
    set_style_font(center_title, east='黑体', west='Times New Roman', size=14, bold=True)
    center_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    center_title.paragraph_format.space_before = Pt(12)
    center_title.paragraph_format.space_after = Pt(12)
    center_title.paragraph_format.line_spacing = Pt(20)
    center_title.paragraph_format.first_line_indent = Pt(0)

    toc_style = get_or_add_style('Thesis TOC')
    set_style_font(toc_style, east='宋体', west='Times New Roman', size=12, bold=False)
    toc_style.paragraph_format.space_before = Pt(1)
    toc_style.paragraph_format.space_after = Pt(1)
    toc_style.paragraph_format.line_spacing = Pt(20)
    toc_style.paragraph_format.first_line_indent = Pt(0)

    caption_style = get_or_add_style('Thesis Caption')
    set_style_font(caption_style, east='宋体', west='Times New Roman', size=10.5, bold=False)
    caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_style.paragraph_format.space_before = Pt(3)
    caption_style.paragraph_format.space_after = Pt(6)
    caption_style.paragraph_format.line_spacing = Pt(20)
    caption_style.paragraph_format.first_line_indent = Pt(0)

    keyword_style = get_or_add_style('Thesis Keyword')
    set_style_font(keyword_style, east='宋体', west='Times New Roman', size=12, bold=False)
    keyword_style.paragraph_format.space_before = Pt(12)
    keyword_style.paragraph_format.space_after = Pt(0)
    keyword_style.paragraph_format.line_spacing = Pt(20)
    keyword_style.paragraph_format.first_line_indent = Pt(0)

    reference_style = get_or_add_style('Thesis Reference')
    set_style_font(reference_style, east='宋体', west='Times New Roman', size=10.5, bold=False)
    reference_style.paragraph_format.space_before = Pt(2)
    reference_style.paragraph_format.space_after = Pt(2)
    reference_style.paragraph_format.line_spacing = Pt(20)
    reference_style.paragraph_format.first_line_indent = Pt(0)


def center_title(text):
    p = doc.add_paragraph(style='Thesis Center Title')
    r = p.add_run(text)
    ef(r, east='黑体', west='Times New Roman', size=14, bold=True)
    return p


def add_auto_toc():
    p = doc.add_paragraph(style='Thesis TOC')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    p.add_run()._r.append(begin)

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\h \\z \\t "Thesis Heading 1,1,Thesis Heading 2,2,Thesis Heading 3,3"'
    p.add_run()._r.append(instr)

    separate = OxmlElement('w:fldChar')
    separate.set(qn('w:fldCharType'), 'separate')
    p.add_run()._r.append(separate)

    hint = p.add_run('目录将随文档打开或字段更新后自动生成')
    ef(hint, east='宋体', west='Times New Roman', size=12, bold=False)

    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    p.add_run()._r.append(end)

    doc.add_paragraph(style='Thesis TOC')

def body(text):
    """正文段落：小四(12pt)宋体，首行缩进2字符(24pt)，固定行距20pt"""
    p = doc.add_paragraph(style='Thesis Body')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(0)
    pf.line_spacing = Pt(20)
    pf.first_line_indent = Pt(24)   # 2×12pt = 24pt
    r = p.add_run(text)
    ef(r, east='宋体', west='Times New Roman', size=12)
    return p

def h1(text):
    """一级标题（章）：四号(14pt)黑体，居左，不加粗"""
    p = doc.add_paragraph(style='Thesis Heading 1')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12); pf.space_after = Pt(12)
    pf.line_spacing = Pt(20)
    pf.first_line_indent = Pt(0)
    pf.left_indent = Pt(0)
    r = p.add_run(text)
    ef(r, east='黑体', size=14, bold=False)
    return p

def h2(text):
    """二级标题（节）：小四(12pt)黑体，左对齐，不加粗"""
    p = doc.add_paragraph(style='Thesis Heading 2')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(6); pf.space_after = Pt(3)
    pf.line_spacing = Pt(20)
    pf.first_line_indent = Pt(0)
    pf.left_indent = Pt(0)
    r = p.add_run(text)
    ef(r, east='黑体', size=12, bold=False)
    return p

def h3(text):
    """三级标题（小节）：小四(12pt)黑体，左对齐，不加粗"""
    p = doc.add_paragraph(style='Thesis Heading 3')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(3); pf.space_after = Pt(0)
    pf.line_spacing = Pt(20)
    pf.first_line_indent = Pt(0)
    pf.left_indent = Pt(0)
    r = p.add_run(text)
    ef(r, east='黑体', size=12, bold=False)
    return p

def insert_fig(fname, caption, w=13.0):
    """插入图片并添加图注（宋体小四居中）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    try:
        run.add_picture(FIGURES + fname, width=Cm(w))
    except Exception:
        ef(run, east='宋体', size=10)
        run.text = f'[图片：{fname}]'
    pc = doc.add_paragraph(style='Thesis Caption')
    pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pc.paragraph_format.space_before = Pt(3)
    pc.paragraph_format.space_after = Pt(6)
    pc.paragraph_format.line_spacing = Pt(20)
    rc = pc.add_run(caption)
    ef(rc, east='宋体', west='Times New Roman', size=10.5)
    return p

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in('w:tcBorders')
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        tag = 'w:' + edge
        element = tcBorders.find(qn(tag))
        if edge_data:
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, value in edge_data.items():
                element.set(qn('w:' + key), str(value))
        elif element is not None:
            tcBorders.remove(element)


def apply_three_line_table(table):
    heavy = {'val': 'single', 'sz': 8, 'color': '000000'}
    thin = {'val': 'single', 'sz': 4, 'color': '000000'}
    last_row = len(table.rows) - 1
    for r, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_border(
                cell,
                top=heavy if r == 0 else None,
                bottom=heavy if r == last_row else (thin if r == 0 else None),
            )


def tbl_add(title, headers, rows):
    """添加表格（含表题）"""
    pt = doc.add_paragraph(style='Thesis Caption')
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pt.paragraph_format.space_before = Pt(6)
    pt.paragraph_format.space_after = Pt(3)
    pt.paragraph_format.line_spacing = Pt(20)
    rt = pt.add_run(title)
    ef(rt, east='宋体', west='Times New Roman', size=10.5)

    # python-docx 的 doc.add_table() 只会按文档尾部插入。
    # 这里先放一个占位段落，再把表移动到表题后面，避免表格串到目录前。
    tail = doc.add_paragraph()
    tail.paragraph_format.space_after = Pt(6)

    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for rn in c.paragraphs[0].runs:
            ef(rn, east='宋体', size=10, bold=True)
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for rn in c.paragraphs[0].runs:
                ef(rn, east='宋体', size=10)
    apply_three_line_table(t)
    pt._p.addnext(t._tbl)

configure_styles()

# ══════════════════════════════════════════════════════════════
# 目 录
# ══════════════════════════════════════════════════════════════
center_title('目  录')
add_auto_toc()

# 中文摘要单独分节，罗马页码从 I 开始
start_section(
    header_text='摘  要',
    show_page_number=True,
    page_number_format='upperRoman',
    page_number_start=1,
)

# ══════════════════════════════════════════════════════════════
# 摘 要
# ══════════════════════════════════════════════════════════════
center_title('摘 要')

body('随着高等教育规模持续扩大，高校选课系统在学期初选课高峰期面临严峻的高并发压力。传统单体架构难以应对大规模并发访问、课程资源动态分配以及个性化推荐等多元需求，系统稳定性、数据一致性和用户体验均存在明显不足。大量高校在选课开放后的前几分钟内即出现服务器响应超时甚至宕机的情况，严重影响学生的正常选课流程，暴露出传统系统架构在突发高并发场景下的脆弱性。')
body('本文设计并实现了一套基于微服务架构的高校智能选课系统，技术选型上以Spring Boot 2.6.13和Spring Cloud Alibaba为核心，Nacos统一管理服务注册与配置，Spring Cloud Gateway对外提供路由入口，Sa-Token处理用户鉴权。选课并发控制方面，引入Redisson分布式锁保护库存扣减临界区，通过RabbitMQ将选课请求异步化从而平滑流量峰值，同时在选课服务内部记录requestId与处理状态，便于结果回查与异常排查。课程推荐模块根据学生专业字段为课程打分，专业对口课程给90分，通用选修课程给50分，降序取前N条作为推荐列表，并附上推荐理由文字；时间冲突检测则在原课程时间窗口基础上两端各延展15分钟，模拟课间转场所需时间，再与已选课程逐一比对。前端选用React结合Ant Design Pro构建管理界面，数据层使用MyBatis-Plus简化了大量重复的CRUD操作。')
body('本文通过接口测试、界面测试和并发压力测试对系统进行了验证。测试结果表明，该系统能够较好地完成登录鉴权、课程推荐、在线选课、消息通知等核心业务流程，并在选课高峰场景下保持较稳定的运行状态，说明所采用的微服务架构、异步消息处理和并发控制方案具有一定的工程应用价值。')

p_kw = doc.add_paragraph()
p_kw.style = 'Thesis Keyword'
rk1 = p_kw.add_run('关键词：')
ef(rk1, east='黑体', size=12, bold=True)
rk2 = p_kw.add_run('微服务；智能推荐；高并发；选课系统；冲突检测')
ef(rk2, east='宋体', west='Times New Roman', size=12)

start_section(
    header_text='Abstract',
    show_page_number=True,
    page_number_format='upperRoman',
    page_number_start=2,
)

# ── ABSTRACT ────────────────────────────────────────────────────
center_title('Abstract')

def ebody(text):
    """英文正文段落"""
    p = doc.add_paragraph(style='Thesis English Body')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.first_line_indent = Pt(24)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)

ebody('With the continuous expansion of higher education, the university course selection system faces severe high-concurrency pressure at the beginning of each semester. Traditional monolithic architectures are inadequate to handle large-scale concurrent access, dynamic course resource allocation, and personalized recommendation requirements, leading to obvious deficiencies in system stability, data consistency, and user experience.')
ebody('This paper designs and implements an intelligent course selection system based on microservice architecture, solving three core problems: system stability and data consistency under high concurrency, personalized intelligent course recommendation, and automatic time-conflict detection and warning. The system uses Spring Boot 2.6.13 with Spring Cloud Alibaba, Nacos for service registration and configuration, Spring Cloud Gateway for unified routing and authentication, and Sa-Token for unified authentication management. Redisson distributed locks protect stock deduction, RabbitMQ handles asynchronous request processing and notifications, and enrollment logs record request status for tracing. The recommendation module scores courses based on major match, assigning 90 points to matched courses and 50 points to general electives. The conflict detection module uses a time-window algorithm with 15-minute buffer thresholds. The frontend is built with React and Ant Design Pro, and the backend data layer uses MyBatis-Plus.')
ebody('The testing results indicate that the system can correctly complete core functions such as authentication, recommendation, course selection, and message notification, while maintaining stable operation under peak course-selection workloads. This demonstrates that the adopted microservice architecture, asynchronous message processing, and concurrency control strategy have practical engineering value.')

p_kw2 = doc.add_paragraph()
p_kw2.style = 'Thesis Keyword'
p_kw2.paragraph_format.first_line_indent = Pt(0)
rk3 = p_kw2.add_run('Key Words: ')
rk3.font.name = 'Times New Roman'; rk3.font.size = Pt(12); rk3.font.bold = True
rk4 = p_kw2.add_run('Microservices; Intelligent Recommendation; High Concurrency; Course Selection; Conflict Detection')
rk4.font.name = 'Times New Roman'; rk4.font.size = Pt(12)

start_section(
    header_text=BODY_HEADER,
    show_page_number=True,
    page_number_format='decimal',
    page_number_start=1,
)
