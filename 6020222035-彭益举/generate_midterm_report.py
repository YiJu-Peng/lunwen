#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""基于中期检查报告模板生成个人成稿。"""

from __future__ import annotations

import re
import subprocess
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_DOC = BASE_DIR / "中期检查报告 (1).doc"
TMP_DIR = BASE_DIR / ".tmp_midterm"
TMP_TEMPLATE_DOCX = TMP_DIR / "中期检查报告 (1).docx"
THESIS_DOCX = BASE_DIR / "2425_41_10475_080902_6020222035_LW.docx"
OUT_DOCX = BASE_DIR / "中期检查报告-彭益举.docx"
OUT_PDF = BASE_DIR / "中期检查报告-彭益举.pdf"


def run(cmd: list[str]) -> None:
    subprocess.run(cmd, check=True)


def convert_template() -> None:
    TMP_DIR.mkdir(exist_ok=True)
    cmd = [
        "libreoffice",
        "--headless",
        "--convert-to",
        "docx",
        "--outdir",
        str(TMP_DIR),
        str(TEMPLATE_DOC),
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0 and not TMP_TEMPLATE_DOCX.exists():
        raise subprocess.CalledProcessError(result.returncode, cmd, output=result.stdout, stderr=result.stderr)


def extract_cover_info() -> dict[str, str]:
    doc = Document(str(THESIS_DOCX))
    info: dict[str, str] = {
        "name": "彭益举",
        "student_id": "6020222035",
        "college": "计算机与信息工程学院",
        "major": "软件工程",
        "class_name": "软件工程2022-1班",
        "advisor": "李佳航",
        "title": "基于微服务架构的高校智能选课系统设计与实现",
    }
    pattern_map = {
        "title": r"题目[:：]\s*(.+)",
        "college": r"学\s*院[:：]\s*(.+)",
        "name": r"姓\s*名[:：]\s*(.+)",
        "student_id": r"学\s*号[:：]\s*(.+)",
        "major": r"专\s*业[:：]\s*(.+)",
        "class_name": r"班\s*级[:：]\s*(.+)",
        "advisor": r"指导教师[:：]\s*(.+)",
    }
    for para in doc.paragraphs[:30]:
        text = para.text.strip()
        if not text:
            continue
        for key, pattern in pattern_map.items():
            match = re.search(pattern, text)
            if match:
                info[key] = match.group(1).strip()
    info["college"] = "软件院"
    return info


def apply_run_style(run, *, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(12)
    if bold is not None:
        run.font.bold = bold


def replace_cell_with_lines(
    cell,
    lines: list[str],
    *,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    bold: bool = False,
) -> None:
    cell.text = ""
    paragraphs = [cell.paragraphs[0]]
    for _ in range(1, len(lines)):
        paragraphs.append(cell.add_paragraph())
    for paragraph, line in zip(paragraphs, lines):
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(line)
        apply_run_style(run, bold=bold)


def replace_cell_simple(cell, text: str, *, align=WD_ALIGN_PARAGRAPH.CENTER, bold: bool = True) -> None:
    replace_cell_with_lines(cell, [text], align=align, bold=bold)


def build_midterm_content(info: dict[str, str]) -> dict[str, list[str] | str]:
    today = date.today()
    today_cn = f"{today.year} 年 {today.month} 月 {today.day} 日"

    honesty_lines = [
        "本人郑重申明：本人所提交的本科毕业论文（设计）是本人在指导教师指导下独立进行研究、写作的成果，论文（设计）中所引用他人的无论以何种方式发布的文字、研究成果，均在论文（设计）中加以说明。",
        "本论文（设计）和资料若有不实之处，本人将承担一切相关责任。",
        f"　　作者签名：  {info['name']}　　　　　　　　　　　　　　　时间： {today_cn}",
    ]

    report_lines = [
        "已完成的研究内容：",
        f"1）围绕“{info['title']}”完成了文献调研、需求分析、总体架构设计、数据库设计和开发环境搭建，明确系统围绕高并发选课、课程推荐、冲突检测和消息通知展开。",
        "2）已完成学生端与管理员端主要功能开发及联调，包括统一登录、课程查询、选课处理、课程表、消息中心、学生管理、课程审核和教师管理等模块。",
        "所取得阶段性成果：",
        "系统已形成从登录鉴权、课程查询到选课处理和结果通知的基本业务闭环，并已完成论文前期章节撰写，整理了系统截图与测试基础材料。",
        "下一步工作计划和研究内容：",
        "继续开展 Postman 或 Apifox 接口测试与 JMeter 压力测试，补充平均响应时间、吞吐量、错误率等量化指标，并同步完成论文后续章节、格式校对和答辩材料整理，形成最终定稿。",
    ]

    teacher_lines = [
        "",
        "指导教师（签名）：",
        "时间：    年     月     日",
    ]

    college_lines = [
        "审查结果：通过（    ）           不通过（    ）",
        "",
        "注：以上结果请在括号内打√",
        "指导教师（签名）：",
        "时           间：",
    ]

    return {
        "gender": "男",
        "topic_change": "（ 无 ）括号内填写有或无。如有，写明原因：",
        "honesty": honesty_lines,
        "report": report_lines,
        "teacher": teacher_lines,
        "college_review": college_lines,
    }


def fill_template() -> None:
    info = extract_cover_info()
    content = build_midterm_content(info)
    doc = Document(str(TMP_TEMPLATE_DOCX))
    table = doc.tables[0]

    replace_cell_simple(table.rows[0].cells[1], info["name"], bold=True)
    replace_cell_simple(table.rows[0].cells[3], content["gender"], bold=True)
    replace_cell_simple(table.rows[0].cells[5], info["student_id"], bold=True)

    replace_cell_simple(table.rows[1].cells[1], info["college"], bold=True)
    replace_cell_simple(table.rows[1].cells[3], info["major"], bold=True)
    replace_cell_simple(table.rows[1].cells[5], info["class_name"], bold=True)

    replace_cell_with_lines(table.rows[2].cells[1], [content["topic_change"]], align=WD_ALIGN_PARAGRAPH.LEFT, bold=True)
    replace_cell_with_lines(table.rows[3].cells[1], [info["title"]], align=WD_ALIGN_PARAGRAPH.CENTER, bold=False)
    replace_cell_with_lines(table.rows[4].cells[1], content["honesty"], align=WD_ALIGN_PARAGRAPH.LEFT, bold=False)
    replace_cell_with_lines(table.rows[5].cells[1], content["report"], align=WD_ALIGN_PARAGRAPH.LEFT, bold=False)
    replace_cell_with_lines(table.rows[6].cells[1], content["teacher"], align=WD_ALIGN_PARAGRAPH.LEFT, bold=False)
    replace_cell_with_lines(table.rows[7].cells[1], content["college_review"], align=WD_ALIGN_PARAGRAPH.LEFT, bold=False)

    doc.save(str(OUT_DOCX))


def convert_pdf() -> None:
    run(
        [
            "libreoffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(BASE_DIR),
            str(OUT_DOCX),
        ]
    )


def main() -> None:
    convert_template()
    fill_template()
    convert_pdf()
    print(f"Done -> {OUT_DOCX}")


if __name__ == "__main__":
    main()
