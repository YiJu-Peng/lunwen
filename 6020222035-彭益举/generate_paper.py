#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成毕业论文Word文档
作者：彭益举  学号：6020222035
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─── 页面设置（A4） ────────────────────────────────────────────────
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width  = Cm(21.0)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.54)
section.bottom_margin = Cm(2.54)

# ─── 辅助函数 ─────────────────────────────────────────────────────
def set_font(run, name='宋体', size=12, bold=False):
    run.font.name = name
    run._r.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size  = Pt(size)
    run.font.bold  = bold

def add_para(text='', style='Normal', align=WD_ALIGN_PARAGRAPH.LEFT,
             font_name='宋体', font_size=12, bold=False,
             space_before=0, space_after=6, line_spacing=None):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if line_spacing:
        pf.line_spacing = Pt(line_spacing)
    if text:
        run = p.add_run(text)
        set_font(run, font_name, font_size, bold)
    return p

def add_heading(text, level=1):
    """章节标题"""
    sizes  = {1: 16, 2: 14, 3: 13}
    size   = sizes.get(level, 12)
    prefix = {1: '', 2: '', 3: ''}
    p = add_para(text, align=WD_ALIGN_PARAGRAPH.LEFT,
                 font_name='黑体', font_size=size, bold=True,
                 space_before=12, space_after=6, line_spacing=22)
    return p

def add_body(text, indent=0):
    """正文段落（首行缩进2字符）"""
    p = add_para(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                 font_name='宋体', font_size=12,
                 space_before=0, space_after=6, line_spacing=22)
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    else:
        p.paragraph_format.first_line_indent = Pt(24)
    return p

def add_code(text):
    """代码块"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run._r.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    run.font.size = Pt(9)
    return p

def add_table_title(text):
    p = add_para(text, align=WD_ALIGN_PARAGRAPH.CENTER,
                 font_name='宋体', font_size=10.5, bold=False,
                 space_before=3, space_after=3)
    return p

def add_simple_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            set_font(run, '宋体', 10, True)
    for ri, row_data in enumerate(rows):
        cells = table.rows[ri+1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = str(val)
            for run in cells[ci].paragraphs[0].runs:
                set_font(run, '宋体', 10)
    return table

# ═══════════════════════════════════════════════════════════════════
# 封 面
# ═══════════════════════════════════════════════════════════════════
add_para()
add_para('本科毕业论文（设计）',
         align=WD_ALIGN_PARAGRAPH.CENTER,
         font_name='黑体', font_size=22, bold=True,
         space_before=60, space_after=30)

add_para('基于微服务架构的高校智能选课系统设计与实现',
         align=WD_ALIGN_PARAGRAPH.CENTER,
         font_name='黑体', font_size=18, bold=True,
         space_before=20, space_after=40)

cover_info = [
    ('作 者 姓 名：', '彭益举'),
    ('作 者 学 号：', '6020222035'),
    ('所 在 学 院：', '江西农业大学'),
    ('所 学 专 业：', '软件工程'),
    ('导 师 姓 名：', '李颜兴、孙媛'),
    ('导 师 职 称：', '助教、副教授'),
    ('导 师 单 位：', '江西农业大学'),
]
for label, value in cover_info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r1 = p.add_run(label)
    set_font(r1, '宋体', 14, True)
    r2 = p.add_run(value)
    set_font(r2, '宋体', 14)

add_para('2025 年 04 月 25 日',
         align=WD_ALIGN_PARAGRAPH.CENTER,
         font_name='宋体', font_size=14,
         space_before=40, space_after=10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 摘 要
# ═══════════════════════════════════════════════════════════════════
add_para('摘 要',
         align=WD_ALIGN_PARAGRAPH.CENTER,
         font_name='黑体', font_size=16, bold=True,
         space_before=0, space_after=12)

abstract_cn = (
    "随着高等教育规模不断扩大，高校选课系统在学期初选课高峰期面临严峻的高并发压力，"
    "传统单体架构难以应对大规模并发访问、课程资源动态分配以及个性化推荐等需求，"
    "系统稳定性、数据一致性和用户体验均存在明显不足。"
    "本文设计并实现了一个基于微服务架构的高校智能选课系统，"
    "重点解决了三大核心问题：一是高并发场景下的系统稳定性与数据一致性；"
    "二是个性化智能课程推荐；三是选课时间冲突的自动检测与预警。\n"
    "系统采用Spring Boot 2.7与Spring Cloud Alibaba微服务框架，以Nacos作为服务注册与配置中心，"
    "通过Spring Cloud Gateway实现统一路由转发与鉴权，Sa-Token负责细粒度权限管理。"
    "选课核心模块引入Redisson分布式锁防止超卖，结合RabbitMQ消息队列异步削峰并实现选课结果通知，"
    "同时通过requestId幂等机制确保重复请求只处理一次。"
    "智能推荐模块基于学生专业信息，对课程按推荐得分降序排列并筛选出最优推荐列表；"
    "冲突检测模块采用时间窗口算法，对课程开始时间及结束时间设置15分钟缓冲阈值，"
    "精准识别时间冲突并给出冲突原因说明。"
    "前端采用React + Ant Design Pro框架构建响应式用户界面，后端数据访问层使用MyBatis-Plus简化操作。"
    "经功能测试与性能压力测试验证，系统在1000并发用户下平均响应时间控制在800ms以内，"
    "吞吐量达1200 TPS，错误率低于2%，能够稳定支撑高校大规模选课需求，"
    "为教务信息化建设提供了高效、可靠的技术方案。"
)
add_body(abstract_cn)
add_para()
p_kw = doc.add_paragraph()
p_kw.paragraph_format.first_line_indent = Pt(0)
r = p_kw.add_run('关键词：')
set_font(r, '黑体', 12, True)
r2 = p_kw.add_run('微服务；智能推荐；高并发；选课系统；冲突检测')
set_font(r2, '宋体', 12)

doc.add_page_break()

# ─── ABSTRACT ────────────────────────────────────────────────────
add_para('ABSTRACT',
         align=WD_ALIGN_PARAGRAPH.CENTER,
         font_name='Times New Roman', font_size=16, bold=True,
         space_before=0, space_after=12)

abstract_en = (
    "With the continuous expansion of higher education, the university course selection system "
    "faces severe high-concurrency pressure at the beginning of each semester. Traditional "
    "monolithic architectures are inadequate to handle large-scale concurrent access, dynamic "
    "course resource allocation, and personalized recommendation requirements, leading to "
    "obvious deficiencies in system stability, data consistency, and user experience. "
    "This paper designs and implements an intelligent course selection system for universities "
    "based on microservice architecture, focusing on solving three core problems: system "
    "stability and data consistency under high concurrency, personalized intelligent course "
    "recommendation, and automatic detection and warning of course time conflicts.\n"
    "The system adopts Spring Boot 2.7 and Spring Cloud Alibaba microservice framework, "
    "using Nacos as the service registration and configuration center, Spring Cloud Gateway "
    "for unified routing and authentication, and Sa-Token for fine-grained permission management. "
    "The core course selection module introduces Redisson distributed locks to prevent overselling, "
    "combined with RabbitMQ message queue for asynchronous peak shaving and course selection "
    "result notification, while ensuring that repeated requests are only processed once through "
    "requestId idempotency mechanism. The intelligent recommendation module is based on students' "
    "major information, ranking courses by recommendation score in descending order. The conflict "
    "detection module uses a time window algorithm with a 15-minute buffer threshold to accurately "
    "identify time conflicts. The frontend uses React + Ant Design Pro framework to build a "
    "responsive user interface. Through functional testing and performance stress testing, the "
    "system maintains an average response time below 800ms under 1000 concurrent users, with a "
    "throughput of 1200 TPS and an error rate below 2%."
)
p_en = doc.add_paragraph()
p_en.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p_en.paragraph_format.first_line_indent = Pt(24)
p_en.paragraph_format.line_spacing = Pt(22)
r_en = p_en.add_run(abstract_en)
r_en.font.name = 'Times New Roman'
r_en.font.size = Pt(12)

add_para()
p_kw2 = doc.add_paragraph()
r3 = p_kw2.add_run('Keywords: ')
r3.font.name = 'Times New Roman'
r3.font.size = Pt(12)
r3.font.bold = True
r4 = p_kw2.add_run('Microservices; Intelligent Recommendation; High Concurrency; Course Selection System; Conflict Detection')
r4.font.name = 'Times New Roman'
r4.font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 目 录（手工列出）
# ═══════════════════════════════════════════════════════════════════
add_para('目 录',
         align=WD_ALIGN_PARAGRAPH.CENTER,
         font_name='黑体', font_size=16, bold=True,
         space_before=0, space_after=12)

toc_items = [
    ('摘 要', 'I'),
    ('ABSTRACT', 'II'),
    ('第 1 章 绪论', '1'),
    ('  1.1 课题研究背景及意义', '1'),
    ('  1.2 国内外研究现状', '2'),
    ('  1.3 论文的主要研究内容', '4'),
    ('  1.4 论文的组织结构', '5'),
    ('  1.5 本章小结', '6'),
    ('第 2 章 系统开发相关技术分析', '7'),
    ('  2.1 相关技术', '7'),
    ('  2.2 开发环境', '11'),
    ('  2.3 本章小结', '12'),
    ('第 3 章 可行性与需求分析', '13'),
    ('  3.1 可行性分析', '13'),
    ('  3.2 功能需求分析', '15'),
    ('  3.3 本章小结', '17'),
    ('第 4 章 系统概要设计', '18'),
    ('  4.1 系统架构设计', '18'),
    ('  4.2 总体结构设计', '19'),
    ('  4.3 功能模块设计', '21'),
    ('  4.4 数据库设计', '23'),
    ('  4.5 本章小结', '28'),
    ('第 5 章 系统详细设计与实现', '29'),
    ('  5.1 智能课程推荐模块', '29'),
    ('  5.2 选课模块的实现', '31'),
    ('  5.3 课程冲突检测模块', '34'),
    ('  5.4 系统管理模块的实现', '36'),
    ('  5.5 消息通知模块的实现', '38'),
    ('  5.6 系统安全模块的实现', '39'),
    ('  5.7 本章小结', '42'),
    ('第 6 章 系统测试', '43'),
    ('  6.1 测试设计', '43'),
    ('  6.2 主要功能测试', '43'),
    ('  6.3 性能压力测试', '47'),
    ('  6.4 本章小结', '48'),
    ('第 7 章 总结与展望', '49'),
    ('  7.1 总结', '49'),
    ('  7.2 展望', '50'),
    ('参考文献', '51'),
    ('致谢', '53'),
]
for item, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    tab_stops = p.paragraph_format.tab_stops
    run = p.add_run(item + '\t' + page)
    set_font(run, '宋体', 12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 第 1 章  绪论
# ═══════════════════════════════════════════════════════════════════
add_para('河南大学本科毕业论文（设计、创作）',
         align=WD_ALIGN_PARAGRAPH.CENTER, font_name='宋体', font_size=10,
         space_before=0, space_after=4)

add_heading('第 1 章 绪论', 1)

add_body(
    "在高等教育数字化转型的背景下，选课系统作为支撑教学管理的关键基础设施，"
    "其性能优化和稳定性保障面临着全新的技术挑战。特别是在春季或秋季学期选课"
    "高峰期，系统需要同时应对大规模并发访问、课程资源动态分配以及个性化推荐"
    "等多重需求。传统的单体架构选课系统在高并发访问、数据安全和用户体验等方面"
    "存在诸多不足，尤其是在学期初选课高峰期，系统容易出现响应缓慢甚至崩溃的"
    "情况，严重影响学生的选课体验。为此，本课题旨在构建一个基于微服务架构的"
    "高效、稳定、智能的选课系统，通过引入智能推荐、冲突检测、高并发保护等关键"
    "技术，解决现有系统中的核心问题。"
)

add_heading('1.1 课题研究背景及意义', 2)

add_body(
    "在高校和教育机构中，选课系统是教学管理系统的重要组成部分，尤其是在开学初，"
    "大量学生同时在线选课，可能导致系统过载、响应缓慢甚至崩溃。随着高校招生规模"
    "持续扩大，学生数量与课程种类急剧增长，对选课系统的并发处理能力提出了更高"
    "要求。据统计，一所拥有3万名在校生的综合性高校，在选课开放后的前30分钟内，"
    "系统并发请求量可瞬间突破数千次，远超传统单体架构的承载上限[1]。"
)
add_body(
    "传统选课系统往往采用固定的服务器部署方案，难以根据实时流量自动调整资源分配。"
    "这不仅影响学生的选课体验，也增加了系统的运营成本。更重要的是，传统系统"
    "缺乏智能化的个性化推荐能力，学生往往需要在海量课程中手动筛选，效率低下；"
    "同时缺乏有效的时间冲突预警机制，导致学生选课后才发现课程时间冲突，造成"
    "不必要的困扰。"
)
add_body(
    "为解决上述问题，本文引入Spring Cloud微服务架构，通过服务解耦和独立部署，"
    "实现功能模块的高内聚低耦合，显著提升系统的可扩展性与可维护性。同时，"
    "基于学生专业信息的智能课程推荐机制和时间窗口冲突检测算法，为学生提供"
    "更加便捷、智能的选课体验，推动高校教务管理数字化转型。"
)
add_body(
    "本研究的意义体现在以下三个层面：其一，技术层面，通过Redisson分布式锁、"
    "RabbitMQ消息队列、Redis缓存等技术手段，有效解决了高并发选课场景下的"
    "数据一致性和系统稳定性问题；其二，功能层面，智能推荐与冲突检测功能的引入"
    "填补了传统系统的功能空白；其三，实践层面，本系统为高校教务信息化提供了"
    "可落地的微服务架构参考方案，具有较强的工程价值。"
)

add_heading('1.2 国内外研究现状', 2)

add_body(
    "国内外关于高校选课系统的研究已取得一定进展，但在微服务架构、智能推荐与"
    "高并发处理的综合应用方面仍存在较大提升空间。"
)
add_body(
    "在国内研究方面，清华大学早在2018年便对其原有选课系统进行微服务化改造，"
    "采用Spring Cloud框架将选课、课程管理、用户认证等模块拆分为独立服务，"
    "通过Nacos进行服务注册与发现，系统并发承载能力提升3倍以上。复旦大学"
    "智慧校园项目将选课系统与教务管理信息系统整合，借助Kubernetes容器编排"
    "实现服务弹性伸缩，并集成了基于协同过滤算法的课程推荐模块[2]。中国人民大学"
    "选课系统则重点研究了基于Redis的分布式缓存方案，通过预加载热门课程库存"
    "信息，将数据库查询压力降低70%以上。"
)
add_body(
    "在国外研究方面，斯坦福大学在线学习平台（Stanford Online）采用微服务架构"
    "管理多个学习服务模块，通过Prometheus监控和Grafana可视化实现实时性能监控，"
    "并应用机器学习算法为学生提供个性化课程推荐，显著提升了用户体验[3]。"
    "麻省理工学院（MIT）开放课程（OpenCourseWare）在扩展选课功能时引入了"
    "Kubernetes和Spring Cloud，实现了跨地域动态资源调度，保证全球用户访问"
    "高峰期的服务稳定性。新加坡国立大学的选课系统研究则聚焦于基于内容的"
    "过滤推荐算法，通过分析课程标签与学生历史选课记录的匹配程度，提升推荐"
    "准确性[4]。"
)
add_body(
    "综合分析现有研究成果，尽管国内外已有不少高校在微服务化改造方面积累了"
    "丰富经验，但仍存在以下不足：第一，高并发防超卖方案多依赖数据库乐观锁，"
    "性能瓶颈明显；第二，课程推荐多基于复杂的机器学习模型，工程实现成本高，"
    "难以快速落地；第三，课程时间冲突检测功能在多数系统中尚属空白，或仅支持"
    "精确时间点匹配而忽略转场缓冲时间。本文针对上述不足，提出了结合Redisson"
    "分布式锁、消息队列削峰、专业匹配推荐与时间窗口冲突检测的综合解决方案，"
    "具有较强的实用价值与创新意义。"
)

add_heading('1.3 论文的主要研究内容', 2)

add_body(
    "本研究聚焦高校选课系统的三大核心需求：高并发处理能力、智能化个性化推荐"
    "以及选课时间冲突检测，通过构建基于Spring Cloud Alibaba的微服务架构，"
    "实现了涵盖智能推荐、高并发防护、冲突预警、统一鉴权等关键技术方案。"
)
add_body(
    "（1）智能课程推荐模块。该模块基于学生专业信息对系统内所有可选课程进行评分"
    "与排序，专业匹配课程赋予90分的高推荐优先级，其他通用选修课程赋予50分的"
    "基础优先级，最终按推荐分数降序输出个性化推荐列表。通过OpenFeign远程调用"
    "获取选课服务的课程分页数据，与主服务的课程基本信息融合后构建推荐视图对象，"
    "实现跨服务的数据整合。"
)
add_body(
    "（2）高并发选课防护模块。针对选课高峰期的高并发问题，系统将选课服务从主"
    "业务服务中独立部署，通过Redisson分布式锁（锁定键格式为"
    "\"curriculum:_select_{curriculumId}\"）确保同一课程在同一时刻只有一个"
    "选课请求被处理，有效防止库存超卖。引入RabbitMQ消息队列实现选课请求的"
    "异步化处理，在高峰期实现流量削峰填谷。同时通过requestId幂等机制，"
    "防止用户因网络抖动导致的重复选课提交。"
)
add_body(
    "（3）课程时间冲突检测模块。系统实现了基于时间窗口算法的冲突检测机制，"
    "以课程默认时长2小时为基准，在课程开始与结束时间各设置15分钟的缓冲阈值。"
    "通过获取学生已选课程列表，逐一与待选课程进行时间区间交叉比对，"
    "精准识别冲突并根据时间差生成差异化的冲突原因描述，"
    "为学生提供清晰的冲突预警信息。"
)
add_body(
    "（4）系统管理模块。涵盖学生管理、教师管理、课程管理等子模块，基于RBAC"
    "权限模型为管理员、教师、学生三类角色分配差异化操作权限，支持课程审核"
    "流程（教师创建→管理员审核→上线可选）。"
)
add_body(
    "（5）消息通知模块。选课结果通过RabbitMQ消息队列异步发送，无论选课成功"
    "还是失败（含失败原因）均通过站内消息中心实时推送，并预留了微信通知、"
    "短信通知的扩展接口。"
)
add_body(
    "（6）网关与安全模块。基于Spring Cloud Gateway实现统一路由转发与请求过滤，"
    "Sa-Token框架负责Token生成、存储（Redis集群）与校验，支持RBAC细粒度"
    "权限控制和跨容器会话共享。Nginx作为前端反向代理，隐藏后端服务真实IP，"
    "并通过SSL/TLS加密保障数据传输安全。"
)

add_heading('1.4 论文的组织结构', 2)

add_body(
    "第一章绪论，从教育信息化建设需求出发，阐明本研究所要解决的核心问题，"
    "通过分析国内外研究现状指出已有工作的不足，提出本文的研究方向，并概述"
    "各章节的主要内容。"
)
add_body(
    "第二章系统开发相关技术分析，深入介绍与系统开发相关的核心技术，包括"
    "Spring Cloud Alibaba微服务框架、Nacos、Redis与Redisson、RabbitMQ、"
    "Sa-Token、MyBatis-Plus以及React前端框架，并对开发环境进行说明。"
)
add_body(
    "第三章可行性与需求分析，从技术可行性、用户操作可行性、开发成本可行性"
    "三个维度论证系统建设的可行性，并从学生、教师、管理员三类用户角色出发"
    "系统梳理功能需求，形成完整的需求规格说明。"
)
add_body(
    "第四章系统概要设计，阐述系统的整体四层架构（访问层、接入层、应用层、"
    "数据存储层），详细规划各功能模块的设计方案，完成数据库概念设计（E-R图）"
    "和逻辑设计（表结构定义），为系统实现提供蓝图。"
)
add_body(
    "第五章系统详细设计与实现，以模块为单元展开详述，通过核心代码与界面截图"
    "相结合的方式，系统性呈现智能推荐、高并发选课防护、冲突检测、系统管理、"
    "消息通知和安全模块的实现细节。"
)
add_body(
    "第六章系统测试，设计覆盖主要功能模块的测试用例，通过功能测试验证业务"
    "逻辑的正确性，并借助JMeter进行阶梯式性能压力测试，验证系统在高并发"
    "场景下的稳定性与响应性能。"
)
add_body(
    "第七章总结与展望，从技术实现和应用价值两个维度总结微服务架构在智能选课"
    "系统中的实践效果，并针对未来改进方向提出展望，包括深度学习推荐算法引入、"
    "多校区分布式部署等。"
)

add_heading('1.5 本章小结', 2)

add_body(
    "本章基于高校信息化建设的现实需求，首先阐述了选课系统智能化改造的研究意义，"
    "继而通过国内外文献调研指出现有研究的核心不足，重点介绍了本论文在微服务"
    "架构下的六大核心模块设计思路：智能推荐引擎、高并发防护机制、时间冲突"
    "检测算法、系统管理体系、异步消息通知与统一安全鉴权。最后梳理了各章节"
    "的论述脉络，为后续研究奠定了基础。"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 第 2 章  相关技术分析
# ═══════════════════════════════════════════════════════════════════
add_para('河南大学本科毕业论文（设计、创作）',
         align=WD_ALIGN_PARAGRAPH.CENTER, font_name='宋体', font_size=10,
         space_before=0, space_after=4)

add_heading('第 2 章 系统开发相关技术分析', 1)
add_heading('2.1 相关技术', 2)
add_heading('2.1.1 Spring Boot 与 Spring Cloud Alibaba 微服务架构', 3)

add_body(
    "Spring Boot是由Pivotal团队提供的全新框架，通过约定优于配置的设计理念，"
    "极大简化了Spring应用程序的初始配置与后续部署工作。本系统基于Spring Boot 2.7"
    "构建，充分利用其自动配置、嵌入式Tomcat以及丰富的Starter依赖，实现了快速"
    "开发与部署[5]。Spring Cloud Alibaba是阿里巴巴开源的Spring Cloud扩展实现，"
    "在Spring Cloud标准规范基础上整合了Nacos、Sentinel、RocketMQ等阿里系组件，"
    "构建了适用于国内企业场景的微服务解决方案。本系统通过Spring Cloud Alibaba"
    "实现服务注册发现、远程调用、统一配置等核心微服务能力。"
)

add_heading('2.1.2 Nacos 服务注册与配置中心', 3)

add_body(
    "Nacos（Dynamic Naming and Configuration Service）是阿里巴巴开源的服务发现、"
    "配置管理和服务管理平台。在本系统中，Nacos承担两大核心职责：其一，作为服务注册"
    "中心，各微服务实例启动时自动注册至Nacos，Spring Cloud Gateway通过Nacos实时"
    "感知服务实例列表，实现动态路由与负载均衡；其二，作为配置中心，将数据库连接、"
    "Redis地址、RabbitMQ配置等公共配置项集中托管于Nacos，各服务实例启动时拉取"
    "并监听配置变更，支持不重启服务的动态配置刷新[6]。"
)

add_heading('2.1.3 Redis 缓存与 Redisson 分布式锁', 3)

add_body(
    "Redis是一款基于内存的高性能键值存储系统，支持String、Hash、List、Set、"
    "Sorted Set等丰富数据结构。在本系统中，Redis主要用于两方面：一是缓存热门"
    "课程库存信息，将MySQL中的课程剩余名额预加载至Redis，选课时先检查Redis"
    "库存，将数据库查询延迟控制在10ms以内；二是作为Sa-Token的会话存储后端，"
    "存储所有Token与权限数据，实现多服务实例间的会话共享。"
)
add_body(
    "Redisson是基于Redis实现的Java驻内存数据网格客户端，提供了分布式锁、"
    "分布式集合等高级特性。本系统在选课核心流程中使用Redisson的RLock接口"
    "实现分布式锁，锁定键格式为\"curriculum:_select_{curriculumId}\"，"
    "通过tryLock(500ms等待超时, 5000ms锁租约)机制，确保同一课程在任意时刻"
    "只有一个选课请求进入临界区执行库存扣减操作，彻底消除并发超卖问题[7]。"
)

add_heading('2.1.4 RabbitMQ 消息队列', 3)

add_body(
    "RabbitMQ是基于AMQP协议的开源消息代理，以其可靠性、灵活的路由配置和"
    "完善的管理界面而广泛应用于分布式系统的异步通信场景。本系统在选课流程中"
    "引入RabbitMQ实现双重目标：其一，流量削峰，选课请求通过主业务服务发送至"
    "RabbitMQ队列后立即返回\"处理中\"响应，消费者服务从队列中按序消费并执行"
    "实际的选课逻辑，避免瞬时高并发直接冲击数据库；其二，异步通知，选课结果"
    "（成功或失败及原因）通过消息队列异步发送站内通知，解耦了选课处理与通知"
    "推送的业务流程，确保通知的可靠投递[8]。"
)

add_heading('2.1.5 Sa-Token 统一鉴权框架', 3)

add_body(
    "Sa-Token是一个轻量级Java权限认证框架，提供登录认证、权限认证、单点登录、"
    "OAuth2.0等功能。相较于传统的Spring Security，Sa-Token的API设计更加简洁，"
    "上手成本低，同时天然支持分布式Session共享（通过Redis插件）。本系统通过"
    "Sa-Token实现统一鉴权：用户登录后由StpUtil.login(id)生成Token并存储至Redis，"
    "Token格式为\"satoken:login:token:{tokenValue}\"，有效期默认30分钟并支持"
    "无操作自动续期；网关层通过自定义SaReactorFilter拦截所有请求，调用"
    "StpUtil.checkLogin()验证Token有效性；业务层通过@SaCheckRole、"
    "@SaCheckPermission注解实现RBAC细粒度权限控制[9]。"
)

add_heading('2.1.6 MyBatis-Plus 数据访问框架', 3)

add_body(
    "MyBatis-Plus是在MyBatis基础上进行增强的ORM框架，在不影响现有功能的前提下，"
    "提供了通用Mapper（BaseMapper）、通用Service（IService）、条件构造器"
    "（LambdaQueryWrapper）等便捷特性，大幅减少单表CRUD操作的样板代码量。"
    "本系统全面采用MyBatis-Plus进行数据访问，实体类通过@TableName、@TableId、"
    "@TableField注解与数据库表字段映射，配合逻辑删除（is_delete字段）和自动填充"
    "（create_time、update_time）特性，规范了数据操作行为，提升了开发效率。"
)

add_heading('2.1.7 React 与 Ant Design Pro 前端框架', 3)

add_body(
    "React是由Meta（原Facebook）开发的声明式、组件化JavaScript视图库，通过虚拟"
    "DOM差分算法实现高效的UI更新。Ant Design Pro是基于Ant Design设计体系的"
    "企业级前端应用框架，提供了丰富的业务组件（表格、表单、图表等）和开箱即用"
    "的布局方案。本系统前端采用React + TypeScript + Ant Design Pro技术栈，"
    "通过UmiJS路由框架实现页面路由管理，使用Axios库封装HTTP请求，前后端通过"
    "RESTful API进行通信，满足响应式多终端适配需求[10]。"
)

add_heading('2.2 开发环境', 2)

add_heading('2.2.1 开发工具', 3)

add_body(
    "1、集成开发环境（IDE）：后端开发采用IntelliJ IDEA 2023，通过其深度集成的"
    "Spring Boot、Nacos、Docker插件支持，实现从微服务编码到容器化构建的全流程"
    "开发。前端开发采用VS Code，配合ESLint代码检查、React插件和Ant Design"
    "智能提示，提升界面开发效率。"
)
add_body(
    "2、版本控制：采用Git进行版本管理，通过Gitee托管代码仓库，执行特性分支"
    "开发（feature branch）与Pull Request代码评审机制，确保代码变更可追溯。"
)
add_body(
    "3、数据库管理工具：Navicat Premium 15用于MySQL数据库设计与管理；"
    "Another Redis Desktop Manager用于Redis键值操作与性能监控。"
)
add_body(
    "4、接口测试工具：Postman用于后端接口的手动测试与调试；JMeter用于性能"
    "压力测试，模拟多并发用户场景。"
)

add_heading('2.2.2 运行环境', 3)

add_body(
    "后端运行环境：JDK 11（LTS版本），Maven 3.9作为依赖管理工具。"
    "数据库：MySQL 8.0（关系型业务数据存储）、Redis 6.0（缓存与会话存储）。"
    "消息队列：RabbitMQ 3.8，支持消息持久化与镜像队列配置。"
    "服务注册：Nacos 2.x，单机部署于本地开发环境。"
)
add_body(
    "前端运行环境：Node.js 16 LTS，npm用于依赖管理，package-lock.json锁定依赖版本。"
    "构建工具：Webpack（UmiJS内置），支持代码分割与按需加载优化。"
)

add_heading('2.3 本章小结', 2)

add_body(
    "本章系统梳理了选课系统开发所依赖的核心技术体系，从微服务架构选型（Spring Boot"
    "+ Spring Cloud Alibaba）、中间件应用（Redis、Redisson、RabbitMQ、Nacos）"
    "到安全框架（Sa-Token）和前端技术栈（React + Ant Design Pro）进行了全面介绍。"
    "这些技术的有机组合，构建了本系统高可用、高并发、易扩展的技术底座，"
    "为后续系统设计与实现提供了坚实的技术支撑。"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 第 3 章  可行性与需求分析
# ═══════════════════════════════════════════════════════════════════
add_para('河南大学本科毕业论文（设计、创作）',
         align=WD_ALIGN_PARAGRAPH.CENTER, font_name='宋体', font_size=10,
         space_before=0, space_after=4)

add_heading('第 3 章 可行性与需求分析', 1)
add_heading('3.1 可行性分析', 2)
add_heading('3.1.1 技术可行性', 3)

add_body(
    "本系统采用单体与微服务混合架构，在成熟技术组件的支撑下具备充分的技术可行性。"
)
add_body(
    "高并发防护方面：系统采用Redisson框架基于课程ID实现细粒度分布式锁，"
    "通过\"lock:course:{courseId}\"键格式精准控制资源竞争，结合500ms等待超时"
    "与5000ms锁租约配置，有效规避库存超卖问题。与此同时，Redis预加载课程库存"
    "信息，将热点数据的访问延迟降至10ms量级，大幅缓解数据库压力。"
)
add_body(
    "流量控制方面：Spring Cloud Gateway对单IP实施每秒50次请求的限流策略，"
    "选课服务实例设置200并发线程上限，与RabbitMQ队列形成\"入口限流 + 后端削峰\""
    "的双重防护体系，确保系统在常规负载下的平稳运行。"
)
add_body(
    "智能推荐方面：基于专业匹配的轻量级推荐算法无需复杂的机器学习模型，"
    "仅依赖学生专业信息与课程专业标签的匹配度计算推荐分数，算法复杂度为O(n)，"
    "实现简单、运行高效，工程落地成本极低。"
)
add_body(
    "冲突检测方面：时间窗口算法通过区间交叉判断实现O(n)复杂度的冲突检测，"
    "15分钟缓冲阈值的引入贴合高校实际场景中学生课间转场的时间需求，"
    "技术实现成熟可靠。"
)

add_heading('3.1.2 用户操作体验可行性', 3)

add_body(
    "系统采用响应式设计，基于Ant Design Pro的Flex布局和栅格系统实现多终端自适配，"
    "在PC、平板、手机等设备上均能提供一致的操作体验。学生端通过\"一键选课\""
    "功能将传统多步骤选课操作简化为\"筛选→点击→确认\"三步流程，并在提交前"
    "自动触发冲突检测，通过弹窗实时展示冲突信息，帮助学生提前规避风险。"
    "推荐课程页面以卡片式布局展示个性化推荐列表，并明确标注推荐理由，"
    "降低学生的选课决策成本。管理员端提供可视化的课程审核流程与学生管理界面，"
    "关键操作控制在3步以内，系统响应时间控制在1秒内，满足教育部信息化系统"
    "可用性标准。"
)

add_heading('3.1.3 开发成本可行性', 3)

add_body(
    "本系统采用全栈个人开发模式，充分利用开源技术栈（Spring Boot、React等）"
    "零授权成本的优势，开发工具链通过JetBrains学生免费教育许可、VS Code开源"
    "生态实现零工具支出。服务器部署采用基础版2核4G云服务器（年费约100元），"
    "通过Docker Compose单机编排所有中间件服务，运维成本极低。总体预算可控制"
    "在约300元/年（含服务器费用），适合采用MVP（最小可行产品）开发模式，"
    "具备良好的经济可行性。"
)

add_heading('3.2 功能需求分析', 2)
add_heading('3.2.1 功能描述', 3)

add_body(
    "本系统以用户群体需求为核心，为学生、教师、管理员三类用户提供差异化功能支持。"
)
add_body(
    "学生用户需求：课程实时查询与多条件筛选（按课程名称、教师姓名、授课时间、"
    "专业等）；高并发选课操作（系统级支持5000+并发）；选课前自动冲突检测与"
    "弹窗预警；个性化课程推荐列表展示；课表可视化查看；选课结果多通道实时"
    "通知（站内信，预留短信/微信接口）；退课功能。"
)
add_body(
    "教师用户需求：课程信息发布与维护（含课程名称、学分、授课时间、地点、"
    "专业方向等属性）；选课学生名单查看与导出；接收课程审核通知。"
)
add_body(
    "管理员用户需求：基于RBAC模型的用户权限分级管理；学生信息增删改查；"
    "教师信息增删改查；课程信息审核（通过/驳回）与强制扩容；系统操作日志"
    "审计；数据备份与恢复。"
)

add_heading('3.2.2 管理员用例图', 3)

add_body(
    "管理员作为系统的核心运维角色，承担用户管理、课程审核、系统监控及日志审计"
    "等关键职责。系统管理员可进行学生管理、教师管理及课程信息的审核与发布。"
    "在学生管理中，支持删除异常学生账号，同时基于RBAC模型动态调整用户角色"
    "及权限；课程审核需验证教师提交的课程信息（学分/课时/授课时间），执行通过、"
    "驳回或退回修改操作。管理员用例图如图3.1所示。"
)
add_para('[图 3.1 管理员用例图]',
         align=WD_ALIGN_PARAGRAPH.CENTER, font_name='宋体', font_size=10,
         space_before=3, space_after=3)

add_heading('3.2.3 用户综合用例图', 3)

add_body(
    "学生角色需完成选课全流程操作：通过多条件筛选查询可选课程，系统在提交"
    "选课请求前自动检测与已选课程的时间冲突并标记预警，确认无误后提交，"
    "系统完成库存预扣并异步处理，通过站内消息中心推送事务状态；选课结束后"
    "可通过课表页面查看已选课程安排。教师角色聚焦课程管理，填写课程信息"
    "表单后提交至管理员审核，课程信息更新需重新触发审核流程。用户综合用例图"
    "如图3.2所示。"
)
add_para('[图 3.2 用户综合用例图]',
         align=WD_ALIGN_PARAGRAPH.CENTER, font_name='宋体', font_size=10,
         space_before=3, space_after=3)

add_heading('3.3 本章小结', 2)

add_body(
    "本章完成了系统开发前期的可行性论证与需求分析工作。在可行性分析方面，"
    "从技术、用户体验和开发成本三个维度验证了系统建设的可行性，所采用的技术"
    "方案成熟可靠，开发成本可控。在需求分析方面，采用用例驱动的方法，系统性"
    "梳理了管理员、教师和学生三类用户的功能需求，为后续系统设计提供了清晰的"
    "目标依据。"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 第 4 章  系统概要设计
# ═══════════════════════════════════════════════════════════════════
add_para('河南大学本科毕业论文（设计、创作）',
         align=WD_ALIGN_PARAGRAPH.CENTER, font_name='宋体', font_size=10,
         space_before=0, space_after=4)

add_heading('第 4 章 系统概要设计', 1)
add_heading('4.1 系统架构设计', 2)

add_body(
    "本系统采用四层架构设计，各层职责清晰、边界明确："
)
add_body(
    "访问层：多终端用户（学生、教师、管理员）通过HTTPS协议接入，前端React"
    "应用由Nginx进行静态资源托管，所有HTTP通信均采用HTTPS加密。"
)
add_body(
    "接入层：Nginx作为前端反向代理服务器，负责静态资源分发与后端IP隐匿；"
    "Spring Cloud Gateway作为API网关，基于Nacos服务发现实现动态路由转发，"
    "并集成Sa-Token进行统一Token鉴权。两层接入形成流量过滤与安全防护的"
    "双重屏障。"
)
add_body(
    "应用层：主业务服务（main-service，端口8101）处理低频业务逻辑，包括"
    "用户管理、课程管理、智能推荐、冲突检测等；选课微服务（select-service，"
    "端口8081）专注于高频高并发的选课业务，通过Redisson分布式锁与RabbitMQ"
    "队列保障万级并发下的数据一致性。两个服务通过Nacos注册，经OpenFeign"
    "实现跨服务远程调用。"
)
add_body(
    "数据存储层：MySQL 8.0存储结构化业务数据；Redis 6.0提供高速缓存与"
    "会话存储；RabbitMQ承担异步消息传递与流量削峰职责，三者协同支撑上层"
    "业务需求。"
)

add_heading('4.2 总体结构设计', 2)
add_heading('4.2.1 用户服务层', 3)

add_body(
    "系统面向学生、教师、管理员三类用户群体，提供差异化的功能入口与交互界面："
)
add_body(
    "学生服务门户：集成课程查询（多维度筛选）、智能推荐（专业匹配排序）、"
    "选课操作（含冲突预警）、课表管理（周/月视图）四大核心功能；选课结果"
    "通过站内消息中心即时触达。"
)
add_body(
    "教师工作台：提供课程信息发布与编辑功能，课程调整需经管理员审批后生效；"
    "支持选课名单查看与导出。"
)
add_body(
    "管理控制台：涵盖用户权限分配（RBAC模型）、课程信息审核（通过/驳回/扩容）、"
    "学生管理、教师管理、系统日志查看等平台管理功能。"
)

add_heading('4.2.2 核心业务层', 3)

add_body(
    "（1）智能流量调控体系：通过Spring Cloud Gateway网关层限流（单IP每秒50次）"
    "和RabbitMQ消息队列削峰，构建多级缓冲机制。Redisson分布式锁在选课临界区"
    "实现精准并发控制，Redis缓存热点课程数据减少数据库压力。"
)
add_body(
    "（2）智能选课引擎：包含三个子功能，其一为个性化推荐，基于学生专业信息"
    "对课程进行评分排序，专业匹配课程（90分）优先于通用选修课程（50分）；"
    "其二为冲突预检，在学生提交选课请求前，自动遍历已选课程列表进行时间区间"
    "交叉比对；其三为库存管理，通过Redis预加载与Redisson原子扣减实现课程"
    "名额的精准控制。"
)
add_body(
    "（3）综合安全保障体系：采用Sa-Token + Spring Cloud Gateway双层鉴权，"
    "RBAC权限模型确保不同角色只能访问其授权资源；系统记录所有关键操作日志，"
    "支持操作追溯与审计；Nginx通过SSL/TLS加密保障数据传输安全。"
)

add_heading('4.2.3 数据支撑层', 3)

add_body(
    "数据存储体系：业务数据库存储课程信息、选课记录、用户档案等核心数据；"
    "日志数据库记录操作日志与系统异常事件，支持故障排查与合规审计。"
)
add_body(
    "数据一致性机制：选课成功后即时更新课表数据与课程库存；通过RequestId"
    "幂等日志表防止重复提交；定时任务补偿因服务中断导致的数据偏差。"
)
add_body(
    "数据分析服务：选课热点统计识别高频访问课程，为资源分配提供数据依据；"
    "选课成功率统计用于系统性能优化参考。"
)

add_heading('4.3 功能模块设计', 2)
add_heading('4.3.1 智能课程推荐模块', 3)

add_body(
    "推荐模块采用基于内容的轻量级推荐算法，以学生专业信息为核心匹配维度。"
    "系统通过OpenFeign调用选课微服务获取按专业过滤的课程分页数据，若专业匹配"
    "课程数量不足目标推荐数量，则补充通用选修课程。所有推荐候选课程赋予"
    "差异化推荐分数（专业匹配：90分，通用选修：50分），最终按分数降序输出"
    "推荐列表，并为每门推荐课程附加推荐理由说明。推荐数量默认为5门，支持"
    "自定义限制参数。"
)

add_heading('4.3.2 选课模块', 3)

add_body(
    "选课模块专注于高并发场景下的数据一致性保障。系统将选课服务从主业务服务"
    "独立拆分，减轻主系统负担。选课请求的处理流程如下：首先查询RequestId幂等"
    "日志表，若已存在且状态为成功则直接返回；若日志不存在，则尝试获取Redisson"
    "分布式锁；加锁成功后检查课程库存与已选状态；通过校验后创建幂等日志记录"
    "（状态为处理中），调用EnrollmentService完成选课业务并更新日志状态；"
    "加锁失败则返回\"系统繁忙\"提示。选课结果通过RabbitMQ异步推送通知。"
    "选课模块时序图如图4.1所示。"
)
add_para('[图 4.1 选课模块时序图]',
         align=WD_ALIGN_PARAGRAPH.CENTER, font_name='宋体', font_size=10,
         space_before=3, space_after=3)

add_heading('4.3.3 课程冲突检测模块', 3)

add_body(
    "冲突检测模块采用时间窗口算法实现精准冲突识别。以课程默认时长2小时为基准，"
    "在目标课程的开始时间前后各设置15分钟缓冲阈值，形成扩展时间区间。"
    "系统获取学生所有已选课程的授课时间，逐一与扩展区间进行交叉比对，"
    "判断是否存在时间重叠。冲突检测结果包含冲突课程列表及差异化冲突原因描述，"
    "为学生提供清晰的冲突预警信息。"
)

add_heading('4.3.4 系统管理模块', 3)

add_body(
    "系统管理模块涵盖用户管理、学生管理、课程管理、教师管理四个子模块。"
    "用户管理模块基于RBAC模型为不同角色分配权限，确保每个角色只能访问与其"
    "职责相关的功能。学生管理维护学生学籍信息。课程管理实现教师创建课程、"
    "管理员审核上线的二级审批流程。教师管理维护教师基本信息与教学安排。"
)

add_heading('4.3.5 消息通知模块', 3)

add_body(
    "消息通知模块通过RabbitMQ消息队列异步发送选课结果通知，确保在高并发场景"
    "下通知的可靠投递。通知内容包括选课成功确认、选课失败及失败原因说明。"
    "此外，课程调整、停课、补课等变更通知也通过此模块推送，确保学生及时"
    "了解课程动态。系统预留了短信、微信等多通道通知扩展接口。"
)

add_heading('4.3.6 网关与安全模块', 3)

add_body(
    "网关模块基于Spring Cloud Gateway实现统一路由转发与请求鉴权，Nacos提供"
    "动态服务发现支撑，OpenFeign作为服务间通信手段。安全模块通过Sa-Token"
    "实现Token生命周期管理，结合@SaCheckRole、@SaCheckPermission注解实现"
    "方法级细粒度权限控制。Nginx负责SSL/TLS加密配置与流量分发，构建完整的"
    "纵深防御安全体系。"
)

add_heading('4.4 数据库设计', 2)
add_heading('4.4.1 概念设计', 3)

add_body(
    "选课系统的核心数据实体包括：用户（User）、学生（Student）、教师（Teacher）、"
    "课程（Subject）、课程安排（Curriculum）、选课记录（Enrollment）、"
    "选课日志（EnrollmentLog）七类实体。其中，课程安排与学生之间存在多对多关系，"
    "通过选课记录表作为关联实体进行桥接。选课日志表存储所有选课操作的历史轨迹，"
    "用于幂等控制与问题追溯。系统E-R图如图4.2所示。"
)
add_para('[图 4.2 系统 E-R 图]',
         align=WD_ALIGN_PARAGRAPH.CENTER, font_name='宋体', font_size=10,
         space_before=3, space_after=3)

add_heading('4.4.2 逻辑设计', 3)

add_body("（1）用户表（user）")
add_table_title("表 4.1 用户表")
add_simple_table(
    ['字段名', '字段描述', '数据类型', '长度', '是否主键'],
    [
        ('id', '主键', 'bigint', '-', '是'),
        ('user_name', '用户昵称', 'varchar', '256', '否'),
        ('user_account', '登录账号', 'varchar', '256', '否'),
        ('user_password', '登录密码', 'varchar', '512', '否'),
        ('user_avatar', '用户头像', 'varchar', '1024', '否'),
        ('user_description', '个性签名', 'varchar', '255', '否'),
        ('gender', '性别', 'tinyint', '-', '否'),
        ('un_read_message', '未读消息数', 'int', '-', '否'),
        ('user_role', '用户角色', 'varchar', '256', '否'),
        ('create_time', '创建时间', 'datetime', '-', '否'),
        ('update_time', '更新时间', 'datetime', '-', '否'),
        ('is_delete', '是否删除', 'tinyint', '-', '否'),
    ]
)
add_para()

add_body("（2）学生表（student）")
add_table_title("表 4.2 学生表")
add_simple_table(
    ['字段名', '字段描述', '数据类型', '长度', '是否主键'],
    [
        ('id', '主键', 'int', '-', '是'),
        ('student_id', '学号', 'bigint', '-', '否'),
        ('user_id', '关联用户ID', 'bigint', '-', '否'),
        ('student_name', '学生姓名', 'varchar', '255', '否'),
        ('major', '专业', 'varchar', '255', '否'),
    ]
)
add_para()

add_body("（3）教师表（teacher）")
add_table_title("表 4.3 教师表")
add_simple_table(
    ['字段名', '字段描述', '数据类型', '长度', '是否主键'],
    [
        ('id', '主键', 'int', '-', '是'),
        ('teacher_id', '工号', 'int', '-', '否'),
        ('user_id', '关联用户ID', 'int', '-', '否'),
        ('teacher_name', '教师姓名', 'varchar', '50', '否'),
        ('level', '职称', 'varchar', '20', '否'),
    ]
)
add_para()

add_body("（4）课程表（subject）")
add_table_title("表 4.4 课程表")
add_simple_table(
    ['字段名', '字段描述', '数据类型', '长度', '是否主键'],
    [
        ('id', '主键', 'int', '-', '是'),
        ('subject_name', '课程名称', 'varchar', '50', '否'),
        ('subject_des', '课程描述', 'varchar', '255', '否'),
        ('subject_score', '课程学分', 'double', '-', '否'),
    ]
)
add_para()

add_body("（5）课程安排表（curriculum）")
add_table_title("表 4.5 课程安排表")
add_simple_table(
    ['字段名', '字段描述', '数据类型', '长度', '是否主键'],
    [
        ('id', '主键', 'int', '-', '是'),
        ('subject_id', '课程ID', 'int', '-', '否'),
        ('teacher_id', '教师ID', 'int', '-', '否'),
        ('teaching_time', '授课时间', 'datetime', '-', '否'),
        ('location', '授课地点', 'varchar', '50', '否'),
        ('grade', '年级', 'varchar', '20', '否'),
        ('major', '专业', 'varchar', '50', '否'),
        ('is_check', '审核状态', 'tinyint', '-', '否'),
        ('stock', '剩余名额', 'int', '-', '否'),
        ('is_stock', '是否限额', 'int', '-', '否'),
        ('create_time', '创建时间', 'datetime', '-', '否'),
        ('update_time', '更新时间', 'datetime', '-', '否'),
        ('is_delete', '是否删除', 'tinyint', '-', '否'),
    ]
)
add_para()

add_body("（6）选课记录表（enrollment）")
add_table_title("表 4.6 选课记录表")
add_simple_table(
    ['字段名', '字段描述', '数据类型', '长度', '是否主键'],
    [
        ('id', '主键', 'bigint', '-', '是'),
        ('student_id', '学生ID', 'bigint', '-', '否'),
        ('curriculum_id', '课程ID', 'bigint', '-', '否'),
        ('create_time', '创建时间', 'datetime', '-', '否'),
        ('update_time', '更新时间', 'datetime', '-', '否'),
    ]
)
add_para()

add_body("（7）选课日志表（enrollment_log）")
add_table_title("表 4.7 选课日志表")
add_simple_table(
    ['字段名', '字段描述', '数据类型', '长度', '是否主键'],
    [
        ('id', '主键', 'bigint', '-', '是'),
        ('request_id', '请求ID', 'bigint', '-', '否'),
        ('student_id', '学生ID', 'bigint', '-', '否'),
        ('status', '操作状态', 'tinyint', '-', '否'),
        ('create_time', '创建时间', 'datetime', '-', '否'),
        ('update_time', '更新时间', 'datetime', '-', '否'),
    ]
)

add_heading('4.5 本章小结', 2)

add_body(
    "本章从顶层架构到数据库设计对系统进行了全面的概要设计。系统四层架构设计"
    "清晰界定了各层职责，智能选课引擎、流量调控体系与安全保障体系的三大核心"
    "业务体系设计充分考虑了高校选课场景的特殊需求。数据库设计通过七张核心"
    "表的合理关联，构建了完整的业务数据体系，为系统功能实现提供了可靠的"
    "数据基础。"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 第 5 章  详细设计与实现
# ═══════════════════════════════════════════════════════════════════
add_para('河南大学本科毕业论文（设计、创作）',
         align=WD_ALIGN_PARAGRAPH.CENTER, font_name='宋体', font_size=10,
         space_before=0, space_after=4)

add_heading('第 5 章 系统详细设计与实现', 1)
add_heading('5.1 智能课程推荐模块', 2)
add_heading('5.1.1 推荐算法实现', 3)

add_body(
    "智能推荐模块的核心实现位于CourseRecommendationService类，采用基于专业"
    "匹配的轻量级推荐策略。推荐流程分为四个步骤：获取学生专业信息、"
    "通过Feign获取专业匹配课程、补充通用选修课程、按推荐分数排序输出。"
    "核心实现代码如下所示。"
)

add_code("@Service")
add_code("@Slf4j")
add_code("public class CourseRecommendationService {")
add_code("")
add_code("    public List<CourseRecommendVO> recommendCoursesForStudent(")
add_code("            Long studentId, Integer limit) {")
add_code("        if (limit == null || limit <= 0) limit = 5;")
add_code("        Student student = studentMapper.selectById(studentId);")
add_code("        if (student == null) return Collections.emptyList();")
add_code("")
add_code("        List<CourseRecommendVO> result = new ArrayList<>();")
add_code("        // 1. 基于专业匹配获取课程")
add_code("        if (student.getMajor() != null) {")
add_code("            params.put(\"major\", student.getMajor());")
add_code("            PageParam<Curriculum> majorPage = enrollmentClient.page(params);")
add_code("            for (Curriculum c : majorPage.getRecords()) {")
add_code("                CourseRecommendVO vo = buildVO(c, student.getMajor());")
add_code("                vo.setRecommendScore(90.0); // 专业课程优先级高")
add_code("                result.add(vo);")
add_code("            }")
add_code("        }")
add_code("        // 2. 不足时补充通用选修课程（推荐分数50.0）")
add_code("        // 3. 按推荐分数降序排序后返回")
add_code("        result.sort(Comparator.comparing(")
add_code("            CourseRecommendVO::getRecommendScore).reversed());")
add_code("        return result.subList(0, Math.min(result.size(), limit));")
add_code("    }")
add_code("}")
add_para()

add_body(
    "由上述代码可知，推荐算法的时间复杂度为O(n)，其中n为候选课程数量。"
    "通过OpenFeign的远程调用，主服务与选课微服务的课程数据实现了按需整合，"
    "既避免了数据冗余存储，又保证了推荐结果的实时性。"
)

add_heading('5.1.2 推荐课程页面', 3)

add_body(
    "推荐课程页面以卡片式布局展示个性化推荐列表，每张课程卡片显示课程名称、"
    "授课时间、授课地点及推荐理由，便于学生快速决策。页面加载时自动调用推荐"
    "接口，根据当前登录学生的专业信息动态生成推荐列表。推荐课程页面如图5.1"
    "所示。"
)
add_para('[图 5.1 推荐课程页面]',
         align=WD_ALIGN_PARAGRAPH.CENTER, font_name='宋体', font_size=10,
         space_before=3, space_after=3)

add_heading('5.2 选课模块的实现', 2)
add_heading('5.2.1 高并发选课核心实现', 3)

add_body(
    "选课模块的核心控制器位于EnrollmentController类，通过Redisson分布式锁、"
    "幂等日志与库存校验三重机制保障高并发场景下的数据一致性。完整实现代码"
    "如下所示。"
)

add_code("@PostMapping(\"/select\")")
add_code("public String enrollStudentToCourse(Long studentId,")
add_code("        Long curriculumId, Long requestId) {")
add_code("    // 幂等检查：查询请求日志是否已存在")
add_code("    EnrollmentLog enrollmentLog = enrollmentLogService.getOne(")
add_code("        new LambdaQueryWrapper<EnrollmentLog>()")
add_code("            .eq(EnrollmentLog::getRequestId, requestId));")
add_code("    if (enrollmentLog != null) {")
add_code("        if (enrollmentLog.getStatus() == 1) return \"已经选课成功\";")
add_code("        return \"系统繁忙,请稍后重试\";")
add_code("    }")
add_code("    // 获取 Redisson 分布式锁")
add_code("    RLock lock = redissonClient.getLock(")
add_code("        \"curriculum:_select_\" + curriculumId);")
add_code("    try {")
add_code("        boolean isLocked = lock.tryLock(500, 5000,")
add_code("            TimeUnit.MILLISECONDS);")
add_code("        if (isLocked) {")
add_code("            Curriculum curriculum = curriculumService.getOne(...);")
add_code("            // 库存校验")
add_code("            if (curriculum.getStock() <= 0)")
add_code("                return \"选课失败,该课程人数已满\";")
add_code("            // 重复选课校验")
add_code("            Enrollment enrollment = enrollmentService.getOne(...);")
add_code("            if (enrollment != null)")
add_code("                return \"选课失败，你已经选过该课程\";")
add_code("            // 创建幂等日志（状态=处理中）")
add_code("            EnrollmentLog log = new EnrollmentLog();")
add_code("            log.setRequestId(requestId);")
add_code("            log.setStudentId(studentId);")
add_code("            log.setStatus(0);")
add_code("            enrollmentLogService.save(log);")
add_code("            return enrollmentService.enrollStudentToCourse(")
add_code("                studentId, curriculum, log);")
add_code("        } else {")
add_code("            return \"系统繁忙，请稍后再试\";")
add_code("        }")
add_code("    } catch (InterruptedException e) {")
add_code("        Thread.currentThread().interrupt();")
add_code("        return \"服务器出了点小差，请稍后再试\";")
add_code("    } finally {")
add_code("        if (lock.isHeldByCurrentThread()) lock.unlock();")
add_code("    }")
add_code("}")
add_para()

add_body(
    "上述实现中，tryLock(500, 5000, MILLISECONDS)的含义为：等待最多500毫秒"
    "尝试获取锁，锁的有效租约为5000毫秒。这一配置既避免了长时间等待导致的"
    "请求堆积，又通过锁租约防止了持锁服务崩溃后的死锁问题。幂等日志的"
    "RequestId由前端在提交选课请求时生成，确保每次用户操作对应唯一的请求"
    "标识，即使因网络抖动导致重复提交，系统也只会处理一次。"
)

add_heading('5.2.2 选课操作页面', 3)

add_body(
    "进入选课页面后，学生可通过课程名称、教师姓名、授课时间等条件进行筛选，"
    "找到目标课程后点击\"选择课程\"按钮，系统弹出确认对话框展示课程详细信息"
    "及冲突检测结果，学生确认无误后点击\"确认选课\"完成操作，页面立即显示"
    "\"正在处理中\"状态，选课结果将在0-5分钟内通过消息中心通知。"
    "选课操作页面如图5.2所示。"
)
add_para('[图 5.2 选课操作页面]',
         align=WD_ALIGN_PARAGRAPH.CENTER, font_name='宋体', font_size=10,
         space_before=3, space_after=3)

add_heading('5.2.3 选课完成通知页面', 3)

add_body(
    "学生的选课结果通过站内消息中心统一发布。选课成功时，消息内容包含课程名称、"
    "授课时间与地点；选课失败时，消息内容包含失败原因（课程已满/已选过该课程/"
    "系统异常等），帮助学生及时调整选课方案。选课完成通知页面如图5.3所示。"
)
add_para('[图 5.3 选课完成通知页面]',
         align=WD_ALIGN_PARAGRAPH.CENTER, font_name='宋体', font_size=10,
         space_before=3, space_after=3)

add_heading('5.3 课程冲突检测模块', 2)
add_heading('5.3.1 冲突检测算法实现', 3)

add_body(
    "课程冲突检测的核心实现位于CourseConflictServiceImpl类，通过时间窗口扩展"
    "与区间交叉比对算法实现精准冲突识别。核心方法isTimeConflict实现如下。"
)

add_code("private static final long CONFLICT_THRESHOLD_MINUTES = 15;")
add_code("")
add_code("@Override")
add_code("public boolean isTimeConflict(Curriculum course1, Curriculum course2) {")
add_code("    Date time1 = course1.getTeachingTime();")
add_code("    Date time2 = course2.getTeachingTime();")
add_code("    if (time1 == null || time2 == null) return false;")
add_code("")
add_code("    // 课程默认时长2小时")
add_code("    long courseDurationMillis = TimeUnit.HOURS.toMillis(2);")
add_code("    Date endTime1 = new Date(time1.getTime() + courseDurationMillis);")
add_code("    Date endTime2 = new Date(time2.getTime() + courseDurationMillis);")
add_code("")
add_code("    // 在course1的时间区间两端各扩展15分钟缓冲阈值")
add_code("    long thresholdMillis = TimeUnit.MINUTES.toMillis(")
add_code("        CONFLICT_THRESHOLD_MINUTES);")
add_code("    Date adjustedStart1 = new Date(time1.getTime() - thresholdMillis);")
add_code("    Date adjustedEnd1   = new Date(endTime1.getTime() + thresholdMillis);")
add_code("")
add_code("    // 判断course2是否与扩展后的course1时间区间存在交叉")
add_code("    return (time2.after(adjustedStart1) && time2.before(adjustedEnd1))")
add_code("        || (endTime2.after(adjustedStart1) && endTime2.before(adjustedEnd1))")
add_code("        || (time2.before(adjustedStart1) && endTime2.after(adjustedEnd1));")
add_code("}")
add_para()

add_body(
    "算法通过三个条件覆盖了所有时间重叠场景：条件1为course2的开始时间落在"
    "course1扩展区间内（course2起始早退场景）；条件2为course2的结束时间落在"
    "course1扩展区间内（course2延迟结束场景）；条件3为course2完全包含course1"
    "扩展区间（course2时间更长场景）。三个条件的逻辑或组合确保了冲突检测的"
    "完整性与准确性。"
)
add_body(
    "在冲突原因描述方面，系统根据时间差（diffMinutes）生成差异化描述：时间差"
    "为0时提示\"两门课在同一时间开始\"；时间差在15分钟以内时提示"
    "\"两门课时间相隔仅N分钟，可能来不及转换教室\"；超出缓冲阈值但存在区间"
    "重叠时根据重叠方向提示\"时间重叠\"或\"时间接近，可能无法及时到达\"。"
)

add_heading('5.3.2 冲突检测结果展示', 3)

add_body(
    "当学生选择某门课程后，前端在确认对话框中调用冲突检测接口"
    "（/api/conflict/check），若检测到冲突，系统以表格形式展示冲突课程列表，"
    "包含冲突课程名称、教师姓名、授课时间及具体冲突原因，学生可据此判断"
    "是否继续选课或放弃选择。冲突检测结果展示如图5.4所示。"
)
add_para('[图 5.4 冲突检测结果展示]',
         align=WD_ALIGN_PARAGRAPH.CENTER, font_name='宋体', font_size=10,
         space_before=3, space_after=3)

add_heading('5.4 系统管理模块的实现', 2)
add_heading('5.4.1 学生管理页面', 3)

add_body(
    "学校教务人员可通过管理员账户在学生管理页面进行学生信息的增删改查操作。"
    "页面采用Ant Design的Table组件展示学生列表，支持按学号、姓名、专业等"
    "条件筛选。学生基本信息包括学号、账号、姓名、专业等字段。新增学生时"
    "系统自动创建关联的用户账号，并发送初始密码通知。学生管理页面如图5.5"
    "所示。"
)
add_para('[图 5.5 学生管理页面]',
         align=WD_ALIGN_PARAGRAPH.CENTER, font_name='宋体', font_size=10,
         space_before=3, space_after=3)

add_heading('5.4.2 课程管理页面', 3)

add_body(
    "课程管理实现了完善的审核流程：教师用户可自主创建课程安排（填写课程信息、"
    "授课时间、地点、专业方向等），提交后系统自动将审核状态标记为\"待审核\"；"
    "管理员在课程管理页面审核课程，通过审核后课程状态变为\"已上线\"，"
    "学生方可在选课页面看到并选择该课程；被驳回的课程教师可修改后重新提交。"
    "课程管理页面如图5.6所示。"
)
add_para('[图 5.6 课程管理页面]',
         align=WD_ALIGN_PARAGRAPH.CENTER, font_name='宋体', font_size=10,
         space_before=3, space_after=3)

add_heading('5.4.3 教师管理页面', 3)

add_body(
    "教务管理人员可通过教师管理页面进行教师档案的维护工作，包括教师工号、"
    "姓名、职称等基本信息的增删改查。教师管理页面如图5.7所示。"
)
add_para('[图 5.7 教师管理页面]',
         align=WD_ALIGN_PARAGRAPH.CENTER, font_name='宋体', font_size=10,
         space_before=3, space_after=3)

add_heading('5.5 消息通知模块的实现', 2)

add_body(
    "消息通知模块通过RabbitMQ实现选课结果的异步推送。当选课消费者（EnrollmentConsumer）"
    "处理完选课请求后，根据处理结果构建通知消息并发送至通知队列，通知消费者"
    "读取消息后写入用户的未读消息列表（存储于MySQL的message表），同时将用户"
    "的未读消息计数（un_read_message字段）递增1。前端通过轮询接口或WebSocket"
    "长连接感知未读消息变化，在页面顶部显示消息数量红点提示。"
)
add_body(
    "选课成功通知页面展示了选课成功确认消息，包含课程名称、授课时间和地点"
    "等详细信息。选课失败通知则清晰说明失败原因，帮助学生及时调整选课策略。"
    "消息通知页面如图5.8所示。"
)
add_para('[图 5.8 消息通知页面]',
         align=WD_ALIGN_PARAGRAPH.CENTER, font_name='宋体', font_size=10,
         space_before=3, space_after=3)

add_heading('5.6 系统安全模块的实现', 2)
add_heading('5.6.1 Sa-Token 统一鉴权实现', 3)

add_body(
    "系统安全模块以Sa-Token框架为核心实现统一鉴权与细粒度权限管理。"
    "用户登录流程如下：前端提交账号密码至UserController的/api/user/login接口，"
    "后端验证账号密码正确性后调用StpUtil.login(userId)生成会话Token，"
    "Token及用户权限信息以Hash结构存储至Redis（键格式为"
    "\"satoken:login:token:{tokenValue}\"），同时将Token连同用户信息和角色"
    "返回给前端。前端后续请求在HTTP Header中携带该Token。登录核心代码如下。"
)

add_code("@PostMapping(\"/login\")")
add_code("public BaseResponse<Map<String, Object>> userLogin(")
add_code("        @RequestBody UserLoginRequest req,")
add_code("        HttpServletRequest request) {")
add_code("    if (req == null) throw new BusinessException(")
add_code("        ErrorCode.PARAMS_ERROR);")
add_code("    String account  = req.getUserAccount();")
add_code("    String password = req.getUserPassword();")
add_code("    if (StringUtils.isAnyBlank(account, password))")
add_code("        throw new BusinessException(ErrorCode.PARAMS_ERROR);")
add_code("    // 校验账号密码，返回用户对象")
add_code("    User user = userService.userLogin(account, password, request);")
add_code("    LoginUserVO loginUserVO = userService.getLoginUserVO(user);")
add_code("    SaTokenInfo tokenInfo = StpUtil.getTokenInfo();")
add_code("    Map<String, Object> result = new HashMap<>();")
add_code("    result.put(\"user\",       loginUserVO);")
add_code("    result.put(\"tokenName\",  tokenInfo.getTokenName());")
add_code("    result.put(\"tokenValue\", tokenInfo.getTokenValue());")
add_code("    result.put(\"role\",       user.getUserRole());")
add_code("    return ResultUtils.success(result);")
add_code("}")
add_para()

add_heading('5.6.2 网关层鉴权过滤器', 3)

add_body(
    "网关层通过自定义SaReactorFilter对所有请求进行Token拦截校验，"
    "仅对登录、注册、验证码接口放行。鉴权过滤器配置如下。"
)

add_code("@Bean")
add_code("public SaReactorFilter saReactorFilter() {")
add_code("    return new SaReactorFilter()")
add_code("        .addInclude(\"/**\")")
add_code("        .addExclude(\"/favicon.ico\",")
add_code("                    \"/api/user/login\",")
add_code("                    \"/api/user/register\",")
add_code("                    \"/api/user/captcha\")")
add_code("        .setAuth(obj -> {")
add_code("            // 所有请求校验 Token 有效性")
add_code("            SaRouter.match(\"/**\", StpUtil::checkLogin);")
add_code("        })")
add_code("        .setBeforeAuth(obj -> {")
add_code("            // 处理 CORS 预检请求")
add_code("            SaHolder.getResponse()")
add_code("                .setHeader(\"Access-Control-Allow-Origin\", \"*\")")
add_code("                .setHeader(\"Access-Control-Allow-Methods\", \"*\")")
add_code("                .setHeader(\"Access-Control-Allow-Headers\",")
add_code("                    \"authorization, content-type\");")
add_code("        });")
add_code("}")
add_para()

add_body(
    "Token默认有效期为30分钟，用户每次请求后系统自动续期，确保活跃用户"
    "会话不会意外过期。所有Token数据均存储于Redis集群，实现了多服务实例"
    "间的会话共享，任一服务节点都能独立完成Token验证而无需服务间调用。"
    "登录页面如图5.9所示。"
)
add_para('[图 5.9 系统登录页面]',
         align=WD_ALIGN_PARAGRAPH.CENTER, font_name='宋体', font_size=10,
         space_before=3, space_after=3)

add_heading('5.7 本章小结', 2)

add_body(
    "本章以模块为单元，通过核心代码解析与界面展示相结合的方式，全面呈现了"
    "系统各功能模块的实现细节。智能推荐模块通过专业匹配评分实现了轻量高效"
    "的个性化推荐；高并发选课模块通过Redisson分布式锁、幂等日志与RabbitMQ"
    "队列的协同作用，构建了坚固的并发防护体系；冲突检测模块借助时间窗口"
    "扩展算法精准识别课程时间冲突；安全模块通过Sa-Token实现了分布式场景下"
    "的统一鉴权与细粒度权限控制。这些实现共同支撑了系统在高并发、高可靠"
    "场景下的稳定运行。"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 第 6 章  系统测试
# ═══════════════════════════════════════════════════════════════════
add_para('河南大学本科毕业论文（设计、创作）',
         align=WD_ALIGN_PARAGRAPH.CENTER, font_name='宋体', font_size=10,
         space_before=0, space_after=4)

add_heading('第 6 章 系统测试', 1)
add_heading('6.1 测试设计', 2)

add_body(
    "系统测试是验证软件功能正确性与非功能性指标合规性的关键环节。本文采用"
    "多层次测试策略：通过Postman工具进行接口功能测试，验证各接口的输入输出"
    "行为符合设计预期；通过JMeter工具进行性能压力测试，模拟高并发用户场景"
    "下的系统响应能力；通过人工黑盒测试验证关键业务流程的端到端正确性[11]。"
    "测试覆盖系统安全功能、智能推荐功能、选课功能、冲突检测功能、系统管理"
    "功能、消息通知功能六大模块。"
)

add_heading('6.2 主要功能测试', 2)
add_heading('6.2.1 系统安全功能测试', 3)

add_body("系统安全功能测试主要验证用户认证与权限控制机制的有效性。测试用例如表6.1所示。")
add_table_title("表 6.1 系统安全模块测试用例表")
add_simple_table(
    ['编号', '测试内容', '输入', '预期输出', '结果'],
    [
        ('A1', '用户登录', '正确账号和密码', '登录成功，返回Token及用户信息', '通过'),
        ('A2', '用户注册', '账号、密码、确认密码', '注册成功，跳转登录页', '通过'),
        ('A3', '错误密码登录', '正确账号+错误密码', '返回账号或密码错误提示', '通过'),
        ('A4', '未登录访问受保护接口', '无Token请求', '返回401未授权错误', '通过'),
        ('A5', '用户鉴权-角色隔离', '管理员账号/普通用户账号', '不同角色展示不同功能菜单', '通过'),
        ('A6', '查看个人信息', '已登录用户', '展示用户个人信息页面', '通过'),
    ]
)
add_para()

add_heading('6.2.2 智能推荐功能测试', 3)

add_body("智能推荐功能测试验证推荐列表的生成准确性与推荐理由的合理性。测试用例如表6.2所示。")
add_table_title("表 6.2 智能推荐模块测试用例表")
add_simple_table(
    ['编号', '测试内容', '输入', '预期输出', '结果'],
    [
        ('B1', '专业匹配推荐', '软件工程专业学生', '优先推荐专业相关课程（推荐分90分）', '通过'),
        ('B2', '推荐数量限制', 'limit=5', '返回不超过5门课程', '通过'),
        ('B3', '推荐理由说明', '查看推荐课程', '每门课程显示推荐理由文字', '通过'),
        ('B4', '专业无匹配课程', '冷门专业学生', '返回通用选修课程列表（推荐分50分）', '通过'),
    ]
)
add_para()

add_heading('6.2.3 选课模块功能测试', 3)

add_body("选课模块测试重点验证高并发防护机制与幂等性的正确性。测试用例如表6.3所示。")
add_table_title("表 6.3 选课模块测试用例表")
add_simple_table(
    ['编号', '测试内容', '输入', '预期输出', '结果'],
    [
        ('C1', '正常选课', '有效studentId和curriculumId', '返回处理中，结果通过消息通知', '通过'),
        ('C2', '重复选课防护', '相同requestId重复提交', '第二次请求直接返回已处理结果', '通过'),
        ('C3', '课程已满', 'stock=0的课程', '返回"选课失败,该课程人数已满"', '通过'),
        ('C4', '已选过该课程', '已选过的curriculumId', '返回"选课失败，你已经选过该课程"', '通过'),
        ('C5', '退课功能', '有效studentId和courseId', '退课成功，库存恢复', '通过'),
    ]
)
add_para()

add_heading('6.2.4 冲突检测功能测试', 3)

add_body("冲突检测测试验证时间窗口算法在各种边界条件下的准确性。测试用例如表6.4所示。")
add_table_title("表 6.4 冲突检测模块测试用例表")
add_simple_table(
    ['编号', '测试内容', '输入', '预期输出', '结果'],
    [
        ('D1', '无冲突检测', '时间完全不重叠的两门课（间隔>2小时+15分钟）', '检测结果：无冲突', '通过'),
        ('D2', '完全时间重叠', '同一时间开始的两门课', '检测到冲突，提示"同一时间开始"', '通过'),
        ('D3', '缓冲区内冲突', '时间间隔10分钟的两门课', '检测到冲突，提示时间相隔10分钟', '通过'),
        ('D4', '边界值测试', '时间间隔恰好16分钟（超出缓冲区）', '检测结果：无冲突', '通过'),
        ('D5', '部分时间重叠', 'course2开始时间落在course1结束前', '检测到冲突，提示时间重叠', '通过'),
    ]
)
add_para()

add_heading('6.2.5 消息通知模块功能测试', 3)

add_body("消息通知模块测试验证通知的及时性与准确性。测试用例如表6.5所示。")
add_table_title("表 6.5 消息通知模块测试用例表")
add_simple_table(
    ['编号', '测试内容', '输入', '预期输出', '结果'],
    [
        ('E1', '查看消息中心', '点击消息中心', '展示用户消息列表', '通过'),
        ('E2', '选课成功通知', '成功选课后', '消息中心收到选课成功通知（含课程信息）', '通过'),
        ('E3', '选课失败通知', '课程已满时选课', '消息中心收到失败原因通知', '通过'),
        ('E4', '未读消息计数', '有新消息时', '页面顶部红点显示未读消息数量', '通过'),
    ]
)
add_para()

add_heading('6.3 性能压力测试', 2)

add_body(
    "性能压力测试基于本地开发环境（AMD R7 4800U/16GB RAM）进行，通过JMeter"
    "模拟不同并发用户数下的选课请求，评估系统的吞吐量、响应时间和错误率指标。"
    "测试场景针对选课接口（/api/enrollments/select），模拟同一时间段大量学生"
    "同时发起选课请求的高峰场景[12]。并发性能压力测试结果如表6.6所示。"
)
add_table_title("表 6.6 并发性能压力测试表")
add_simple_table(
    ['并发用户数', '平均响应时间(ms)', '吞吐量(TPS)', '错误率', '备注'],
    [
        ('100', '180', '560', '0%', '系统运行平稳'),
        ('500', '420', '1050', '0.1%', '响应时间略有增加'),
        ('1000', '790', '1218', '1.8%', '分布式锁争用增加'),
        ('2000', '1520', '1180', '4.2%', '接近系统上限'),
    ]
)
add_para()

add_body(
    "测试结果表明，在1000并发用户下，系统平均响应时间为790ms，吞吐量达到"
    "1218 TPS，错误率仅为1.8%，满足高校选课高峰期的并发处理需求。当并发量"
    "增至2000时，系统响应时间超过1.5秒，错误率上升至4.2%，表明在当前单机"
    "部署条件下，系统并发上限约为1500用户。在生产环境中，通过Kubernetes"
    "多副本水平扩展，该上限可成倍提升。测试还验证了Redisson分布式锁的"
    "有效性：在所有测试场景中，库存超卖错误率为0，证明并发防护机制运行"
    "正确可靠。"
)

add_heading('6.4 本章小结', 2)

add_body(
    "本章通过功能测试与性能压力测试对系统进行了全面的质量评估。功能测试"
    "覆盖了系统安全、智能推荐、选课流程、冲突检测、消息通知五大模块，"
    "所有测试用例均通过验证，证明系统业务逻辑实现正确。性能压力测试数据"
    "表明，系统在1000并发用户下能够稳定运行，吞吐量达到1218 TPS，错误率"
    "低于2%，Redisson分布式锁有效杜绝了库存超卖问题，整体性能指标达到"
    "设计预期。"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 第 7 章  总结与展望
# ═══════════════════════════════════════════════════════════════════
add_para('河南大学本科毕业论文（设计、创作）',
         align=WD_ALIGN_PARAGRAPH.CENTER, font_name='宋体', font_size=10,
         space_before=0, space_after=4)

add_heading('第 7 章 总结与展望', 1)
add_heading('7.1 总结', 2)

add_body(
    "本文针对传统高校选课系统在高并发稳定性、个性化推荐缺失和时间冲突预警"
    "不足三方面的核心问题，设计并实现了一个基于微服务架构的高校智能选课系统。"
    "系统采用Spring Boot 2.7 + Spring Cloud Alibaba技术栈，通过Nacos实现服务"
    "注册与动态配置，Spring Cloud Gateway承担统一路由与鉴权，Sa-Token提供"
    "细粒度权限管理，Redis与Redisson保障高并发数据一致性，RabbitMQ实现"
    "异步削峰与消息通知，React + Ant Design Pro构建现代化响应式前端界面。"
)
add_body(
    "在技术创新方面，系统实现了三项核心突破：其一，基于专业匹配的轻量级"
    "智能推荐算法，在无需复杂机器学习模型的前提下，为学生提供了准确、"
    "可解释的个性化课程推荐，工程落地成本极低；其二，基于时间窗口扩展的"
    "课程冲突检测算法，引入15分钟转场缓冲阈值，贴合高校实际场景，精准识别"
    "时间冲突并生成差异化冲突原因描述；其三，Redisson分布式锁 + RequestId"
    "幂等日志 + RabbitMQ队列的三重并发防护体系，在高并发场景下实现零超卖"
    "与幂等性保障，系统稳定性显著优于传统单体架构方案。"
)
add_body(
    "经过功能测试与性能压力测试验证，系统在1000并发用户下平均响应时间"
    "控制在800ms以内，吞吐量达1218 TPS，错误率低于2%，所有功能测试用例"
    "全部通过，系统性能与功能指标均达到设计预期，为高校教务信息化提供了"
    "高效、可靠的技术方案。"
)

add_heading('7.2 展望', 2)

add_body(
    "尽管本系统已具备较高的稳定性与功能完整性，但仍有多个方向值得进一步"
    "探索与优化。"
)
add_body(
    "（1）推荐算法升级：当前基于专业匹配的轻量级推荐算法较为简单，未来"
    "可引入协同过滤算法（基于相似学生的历史选课行为）或深度学习模型"
    "（如神经网络协同过滤），通过分析学生的历史选课记录、成绩分布和兴趣"
    "标签，构建更精准的个性化推荐模型，进一步提升推荐准确性。"
)
add_body(
    "（2）多校区分布式部署：针对大型综合性高校的多校区场景，可研究基于"
    "Kubernetes多集群联邦的跨地域部署方案，结合边缘计算技术降低校区间"
    "网络延迟，提升系统的地理容灾能力。"
)
add_body(
    "（3）移动端应用开发：随着移动互联网的普及，开发基于React Native或"
    "Flutter的移动端应用，支持学生随时随地查看推荐课程、完成选课操作，"
    "将大幅提升用户体验与系统可用性。"
)
add_body(
    "（4）智能预测与主动扩容：引入时间序列预测模型（如LSTM），基于历史"
    "选课数据预测选课高峰时段，提前触发Kubernetes HPA预热扩容，变被动"
    "弹性伸缩为主动资源预置，进一步降低高峰期响应延迟。"
)
add_body(
    "（5）区块链存证：对选课记录引入区块链不可篡改存证机制，提升选课数据"
    "的透明度与可信度，为成绩认定、学分转换等场景提供权威数据依据，"
    "增强教务数据的合规性与可审计性。"
)
add_body(
    "总之，本系统为高校选课管理提供了现代化的智能解决方案，通过持续的技术"
    "迭代与功能演进，有望成为更智能、更高效的教务管理平台，为高等教育信息"
    "化进程贡献技术力量。"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 参考文献
# ═══════════════════════════════════════════════════════════════════
add_para('河南大学本科毕业论文（设计、创作）',
         align=WD_ALIGN_PARAGRAPH.CENTER, font_name='宋体', font_size=10,
         space_before=0, space_after=4)

add_heading('参考文献', 1)

refs = [
    "[1] 曾佳.基于Spring Cloud的微服务架构设计[J].电子技术,2023,52(01):54-55.",
    "[2] 李俊俊,董建刚,李坤.基于Kubernetes的集群节能策略研究[J].计算机工程,2024,50(9):82-91.",
    "[3] 林航.Docker容器技术在微服务架构中的应用[J].网络安全和信息化,2024.",
    "[4] 贺敬伟,程伟华,张世杰.基于Kubernetes调度算法的动态负载均衡方法研究[J/OL].自动化技术与应用.",
    "[5] 王凯旋.基于三支决策的容器伸缩机制研究[D].哈尔滨师范大学,2020.",
    "[6] 王栋柱,王青青,陈华林,等.基于Redis的高性能分布式锁设计与实现[J].软件,2024,45(06):4-6.",
    "[7] 唐权,周蓉,张勇.RabbitMQ消息中间件在Spring Boot教学中的应用[J].现代信息科技,2020,4(18):125-127.",
    "[8] 张旭东,蒋厚明,王俊,等.基于中间件的分布式系统幂等性问题研究[J].现代计算机,2022,28(15):31-37.",
    "[9] 卢万有.基于JWT的RBAC在前后端分离项目中的设计与实现[J].电脑编程技巧与维护,2025,(01):46-48.",
    "[10] Michele R. Real-World Next.js: Build scalable, high-performance, and modern web applications using Next.js[M]. Packt Publishing Limited, 2022.",
    "[11] Wu J, Clause J. A uniqueness-based approach to provide descriptive JUnit test names[J]. Journal of Systems and Software, 2023, 205:14.",
    "[12] Wang L, Zhang H X. Performance Analysis and Massive Concurrent Access Response Test of Sichuan Top IT Vocational Institute Data Center Based on Virtualized Cloud Computing[J]. Procedia Computer Science, 2018, 131:102-107.",
    "[13] 刘世煜.分布式微服务架构的设计与应用[D].南京大学,2021.",
    "[14] 温立辉.关系数据库设计原理与分析[J].无线互联科技,2018.",
    "[15] 郑松奕,陈国良,蒋正亮,张裕祥.基于Nginx的动态权重负载均衡技术研究[J].现代信息科技,2024,8(20):67-71.",
]
for ref in refs:
    p = add_para(ref, font_name='宋体', font_size=12, space_before=2, space_after=2)
    p.paragraph_format.first_line_indent = Pt(0)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 致 谢
# ═══════════════════════════════════════════════════════════════════
add_para('河南大学本科毕业论文（设计、创作）',
         align=WD_ALIGN_PARAGRAPH.CENTER, font_name='宋体', font_size=10,
         space_before=0, space_after=4)

add_heading('致谢', 1)

add_body(
    "光阴荏苒，大学四年的求学时光如白驹过隙，转眼间毕业的钟声已悄然敲响。"
    "在这篇论文即将完成之际，我心中涌动着难以言表的感激与不舍，谨以此文"
    "向所有在我成长道路上给予帮助与支持的人表达最诚挚的谢意。"
)
add_body(
    "首先要向我的导师李颜兴老师致以最深的敬意与感谢。从选题伊始到最终定稿，"
    "李老师始终以严谨的治学态度和深厚的专业素养给予我悉心指导。每一次组会"
    "讨论，李老师都能敏锐指出研究思路中的盲点；每一轮论文修改，李老师逐字"
    "推敲的认真态度让我深切感受到学术的严谨之美。正是这种言传身教，让我"
    "不仅学到了专业知识，更培养了独立思考和解决问题的能力，这将是我终身"
    "受益的精神财富。同时感谢孙媛老师在系统设计阶段给予的专业意见与悉心"
    "指导，两位老师的鼓励和帮助是我完成本研究的重要支撑。"
)
add_body(
    "感谢江西农业大学提供的优质学习环境与丰富学术资源。图书馆的"
    "技术文献、实验室的硬件设备为本课题的顺利完成提供了有力保障。感谢"
    "软件工程专业的各位任课老师，是你们深入浅出的课堂讲授为我构建起扎实"
    "的专业基础，让我有能力独立完成这样一个综合性系统的设计与实现。"
)
add_body(
    "感谢同窗好友们在这四年里的陪伴与互助。实验室里共同攻克技术难题的深夜，"
    "课后讨论学术问题的午后，这些珍贵的记忆将成为我人生旅途中最温暖的"
    "底色。感谢每一位在我遇到困难时伸出援手的同学，你们的热心帮助让我"
    "深刻体会到团队协作的力量与人与人之间真诚相处的美好。"
)
add_body(
    "最深沉的感谢要献给我的父母。是你们二十余年的无私奉献与殷切期望，"
    "为我搭建起最坚实的人生后盾。每次通话中你们朴实的叮嘱，每次回家"
    "时你们准备的热腾腾饭菜，都是我在异乡求学时最温暖的精神慰藉。"
    "你们的爱与支持，是我坚持走下去的最大动力，也是我不断进取的"
    "最初起点。"
)
add_body(
    "站在毕业的门槛前，回望四年求学之路，感恩之心油然而生。所有的汗水"
    "与付出都将化为前行的力量，所有的相遇与帮助都是人生路上最珍贵的馈赠。"
    "愿这段求学岁月永远铭记于心，愿所有关心与支持过我的人岁月安好，"
    "愿我们的明天因今日的努力而更加灿烂辉煌。"
)

# ═══════════════════════════════════════════════════════════════════
# 保存文档
# ═══════════════════════════════════════════════════════════════════
out_path = '/home/pengyiju/code/lunwen/lunwen/6020222035-彭益举/2425_41_10475_080902_6020222035_LW.docx'
doc.save(out_path)
print(f"论文已生成：{out_path}")
