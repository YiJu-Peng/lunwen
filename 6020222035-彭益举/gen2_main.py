#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""主入口：组合 part1 + part2 + part3 生成完整论文并刷新目录"""
import sys
sys.path.insert(0, '/home/pengyiju/code/lunwen/lunwen/6020222035-彭益举/')

import importlib.util, types
import os
import re
import socket
import subprocess
import time
from pathlib import Path

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

def load(path):
    spec = importlib.util.spec_from_file_location('m', path)
    m = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(m)
    return m


def pick_free_port():
    try:
        with socket.socket() as sock:
            sock.bind(('127.0.0.1', 0))
            return sock.getsockname()[1]
    except PermissionError:
        return None


def wait_for_port(port, timeout=15):
    deadline = time.time() + timeout
    while time.time() < deadline:
        with socket.socket() as sock:
            sock.settimeout(0.5)
            if sock.connect_ex(('127.0.0.1', port)) == 0:
                return True
        time.sleep(0.2)
    return False


def extract_toc_entries(doc_path):
    """用 LibreOffice 计算目录页码，再把目录条目提取出来。"""
    doc_path = Path(doc_path).resolve()
    try:
        import uno
        from com.sun.star.beans import PropertyValue
    except Exception as exc:
        print(f'Skip TOC extraction: UNO unavailable ({exc})')
        return []

    port = pick_free_port()
    if port is None:
        print('Skip TOC extraction: socket permission denied in current environment')
        return []
    accept = f'socket,host=127.0.0.1,port={port};urp;StarOffice.ServiceManager'
    proc = subprocess.Popen(
        [
            'libreoffice',
            '--headless',
            '--nologo',
            '--nodefault',
            '--nofirststartwizard',
            '--nolockcheck',
            f'--accept={accept}',
        ],
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )
    try:
        if not wait_for_port(port):
            print('Skip field refresh: LibreOffice listener did not start in time')
            return

        local_ctx = uno.getComponentContext()
        resolver = local_ctx.ServiceManager.createInstanceWithContext(
            'com.sun.star.bridge.UnoUrlResolver',
            local_ctx,
        )
        ctx = resolver.resolve(
            f'uno:socket,host=127.0.0.1,port={port};urp;StarOffice.ComponentContext'
        )
        desktop = ctx.ServiceManager.createInstanceWithContext(
            'com.sun.star.frame.Desktop',
            ctx,
        )

        hidden = PropertyValue()
        hidden.Name = 'Hidden'
        hidden.Value = True

        update_mode = PropertyValue()
        update_mode.Name = 'UpdateDocMode'
        update_mode.Value = 1

        model = desktop.loadComponentFromURL(
            uno.systemPathToFileUrl(str(doc_path)),
            '_blank',
            0,
            (hidden, update_mode),
        )
        indexes = model.getDocumentIndexes()
        for idx in range(indexes.getCount()):
            indexes.getByIndex(idx).update()
        entries = []
        enum = model.Text.createEnumeration()
        seen_toc_title = False
        while enum.hasMoreElements():
            para = enum.nextElement()
            try:
                text = para.getString().strip()
            except Exception:
                text = ''
            if not text:
                continue
            if text == '目  录':
                seen_toc_title = True
                continue
            if seen_toc_title:
                if '\t' in text:
                    title, page = text.rsplit('\t', 1)
                    entries.append((title.strip(), page.strip()))
                    continue
                if entries:
                    break
        model.close(True)
        return entries
    except Exception as exc:
        print(f'Skip TOC extraction: {exc}')
        return []
    finally:
        proc.terminate()
        try:
            proc.wait(timeout=5)
        except subprocess.TimeoutExpired:
            proc.kill()


def get_toc_level(title):
    if re.match(r'^\d+\.\d+\.\d+', title):
        return 3
    if re.match(r'^\d+\.\d+', title):
        return 2
    return 1


def normalize_docx_layout(doc_path):
    """通过一次 python-docx round-trip 减少 LibreOffice 的分页异常。"""
    doc_path = Path(doc_path).resolve()
    doc = Document(str(doc_path))
    doc.save(str(doc_path))


def build_static_toc(doc_path, entries):
    if not entries:
        return False

    doc_path = Path(doc_path).resolve()
    doc = Document(str(doc_path))
    paragraphs = doc.paragraphs

    toc_title_idx = None
    for idx, p in enumerate(paragraphs):
        if p.text.strip() == '目  录':
            toc_title_idx = idx
            break
    if toc_title_idx is None:
        print('Skip TOC rewrite: TOC title not found')
        return False

    section_para = None
    for p in paragraphs[toc_title_idx + 1:]:
        if p._p.xpath('./w:pPr/w:sectPr'):
            section_para = p
            break
    if section_para is None:
        print('Skip TOC rewrite: section break after TOC not found')
        return False

    node = doc.paragraphs[toc_title_idx]._p.getnext()
    while node is not None and node is not section_para._p:
        nxt = node.getnext()
        node.getparent().remove(node)
        node = nxt

    right_tab = Cm(14.8)
    for title, page in entries:
        level = get_toc_level(title)
        p = doc.add_paragraph(style='Thesis TOC')
        pf = p.paragraph_format
        pf.left_indent = Pt(12 * (level - 1))
        pf.first_line_indent = Pt(0)
        pf.space_before = Pt(1)
        pf.space_after = Pt(1)
        pf.line_spacing = Pt(20)
        pf.tab_stops.add_tab_stop(right_tab, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

        run = p.add_run(f'{title}\t{page}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = False
        run._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '宋体')

        section_para._p.addprevious(p._p)

    doc.save(str(doc_path))
    return True

# ── 执行 Part1（生成 doc 对象 + 封面/摘要/目录）────────────────
p1 = load('/home/pengyiju/code/lunwen/lunwen/6020222035-彭益举/gen2_part1.py')
doc       = p1.doc
body      = p1.body
h1        = p1.h1
h2        = p1.h2
h3        = p1.h3
insert_fig = p1.insert_fig
tbl_add   = p1.tbl_add
start_section = p1.start_section
OUT       = p1.OUT

# ── 执行 Part2（第1-3章）────────────────────────────────────────
p2 = load('/home/pengyiju/code/lunwen/lunwen/6020222035-彭益举/gen2_part2.py')
p2.write_ch1_3(doc, body, h1, h2, h3, insert_fig, tbl_add)

# ── 执行 Part3（第4-7章 + 参考文献 + 致谢）──────────────────────
p3 = load('/home/pengyiju/code/lunwen/lunwen/6020222035-彭益举/gen2_part3.py')
p3.write_ch4_7(doc, body, h1, h2, h3, insert_fig, tbl_add, start_section)

# ── 保存 ────────────────────────────────────────────────────────
doc.save(OUT)
normalize_docx_layout(OUT)
if os.environ.get('THESIS_STATIC_TOC', '1') != '0':
    toc_entries = extract_toc_entries(OUT)
    build_static_toc(OUT, toc_entries)
print(f'Done -> {OUT}')
